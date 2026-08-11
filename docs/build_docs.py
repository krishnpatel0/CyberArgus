"""Build Argus Unified implementation documentation and monochrome flowcharts.

The reports intentionally describe only behavior present in the consolidated
repository.  Existing source-project DOCX files were used as lineage material;
planned, aspirational, or obsolete architecture from those documents is not
carried forward as implemented behavior.
"""

from __future__ import annotations

import math
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from lxml import etree
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MODULES = DOCS / "modules"
FLOWS = DOCS / "flowcharts"
ASSETS = DOCS / "assets"
RENDERS = DOCS / "renders"
for directory in (MODULES, FLOWS, ASSETS, RENDERS):
    directory.mkdir(parents=True, exist_ok=True)

DATE = "10 August 2026"
VERSION = "Unified implementation report · revision 1.0"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
MID_GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
BLACK = "000000"
WHITE = "FFFFFF"
CONTENT_DXA = 9360  # 6.5 in at 1440 DXA/in


SOURCES = {
    "opencti": ("OpenCTI documentation", "https://docs.opencti.io/latest/usage/getting-started/"),
    "opencti_connectors": ("OpenCTI connector documentation", "https://docs.opencti.io/latest/deployment/connectors/"),
    "misp": ("MISP feature overview", "https://www.misp-project.org/features/"),
    "threatconnect_indicators": ("ThreatConnect indicator API", "https://docs.threatconnect.com/en/latest/rest_api/v3/indicators/indicators.html"),
    "threatconnect_associations": ("ThreatConnect associations API", "https://docs.threatconnect.com/en/latest/rest_api/v3/associations.html"),
    "threatconnect_cases": ("ThreatConnect case-management SDK", "https://docs.threatconnect.com/en/latest/tcex/module_case_management.html"),
    "threatconnect_endpoints": ("ThreatConnect API endpoint catalog", "https://docs.threatconnect.com/en/latest/rest_api/v3/available_endpoints.html"),
    "threatconnect_roles": ("ThreatConnect system roles", "https://docs.threatconnect.com/en/latest/rest_api/v3/system_roles/system_roles.html"),
    "threatconnect_stix": ("ThreatConnect STIX support release notes", "https://docs.threatconnect.com/en/latest/tcex/release_notes.html"),
    "recorded_future": ("Recorded Future Intelligence Cloud", "https://www.recordedfuture.com/platform"),
    "crowdstrike": ("CrowdStrike adversary intelligence", "https://www.crowdstrike.com/en-us/platform/threat-intelligence/adversary-intelligence/"),
    "spiderfoot": ("SpiderFoot official repository", "https://github.com/smicallef/spiderfoot"),
    "maltego": ("Maltego Transform Hub", "https://www.maltego.com/transform-hub/"),
    "maltego_transforms": ("Maltego transform execution", "https://docs.maltego.com/en/support/solutions/articles/15000009613-running-transforms"),
    "maltego_collaboration": ("Maltego graph collaboration", "https://docs.maltego.com/en/support/solutions/articles/15000010791-collaboration"),
    "sherlock": ("Sherlock official repository", "https://github.com/sherlock-project/sherlock"),
    "maigret": ("Maigret official repository", "https://github.com/soxoj/maigret"),
    "hibp": ("Have I Been Pwned API v3", "https://haveibeenpwned.com/scalar/"),
    "intelx": ("Intelligence X API", "https://help.intelx.io/api/"),
    "leakcheck": ("LeakCheck API documentation", "https://docs.leakcheck.io/overview"),
    "dehashed": ("DeHashed", "https://www.dehashed.com/"),
    "telethon": ("Telethon quick start", "https://docs.telethon.dev/en/stable/basic/quick-start.html"),
    "telethon_events": ("Telethon update events", "https://docs.telethon.dev/en/stable/basic/updates.html"),
    "maltego_monitor": ("Maltego Monitor", "https://www.maltego.com/monitor/"),
    "maltego_telegram": ("Maltego Telegram transforms", "https://docs.maltego.com/en/support/solutions/articles/15000058961"),
    "sociallinks": ("Social Links platform capabilities", "https://sociallinks.io/"),
    "sociallinks_telegram": ("Social Links Telegram Pack", "https://blog.sociallinks.io/social-links-product-updates-q2-2023/"),
    "presidio": ("Microsoft Presidio", "https://microsoft.github.io/presidio/"),
    "presidio_entities": ("Presidio supported PII entities", "https://microsoft.github.io/presidio/supported_entities/"),
    "presidio_anonymizer": ("Presidio anonymizer", "https://microsoft.github.io/presidio/anonymizer/"),
    "stix": ("OASIS STIX 2.1 introduction", "https://oasis-open.github.io/cti-documentation/stix/intro.html"),
    "google_dlp": ("Google Sensitive Data Protection text inspection", "https://docs.cloud.google.com/sensitive-data-protection/docs/inspecting-text"),
    "google_dlp_custom": ("Google custom infoType detectors", "https://docs.cloud.google.com/sensitive-data-protection/docs/creating-custom-infotypes"),
    "macie": ("Amazon Macie managed identifiers", "https://docs.aws.amazon.com/macie/latest/user/managed-data-identifiers.html"),
    "macie_custom": ("Amazon Macie custom identifiers", "https://docs.aws.amazon.com/macie/latest/user/custom-data-identifiers.html"),
    "iocextract": ("iocextract official repository", "https://github.com/pedramamini/iocextract"),
    "nist_risk": ("NIST SP 800-30 Rev. 1 — Guide for Conducting Risk Assessments", "https://doi.org/10.6028/NIST.SP.800-30r1"),
    "temporal_ir": ("Campos et al. — Temporal Information Retrieval", "https://doi.org/10.1561/1500000043"),
    "fellegi_sunter": ("Fellegi and Sunter — A Theory for Record Linkage", "https://doi.org/10.1080/01621459.1969.10501049"),
    "jaccard": ("Jaccard — Étude comparative de la distribution florale", "https://doi.org/10.5169/seals-266450"),
    "meta_blocking": ("Papadakis et al. — Enhanced Meta-Blocking for Entity Resolution", "https://doi.org/10.1016/j.bdr.2016.08.002"),
    "moore_bfs": ("Moore — The Shortest Path Through a Maze", "https://books.google.com/books?id=IVZBHAAACAAJ"),
    "fips180": ("NIST FIPS 180-4 — Secure Hash Standard", "https://doi.org/10.6028/NIST.FIPS.180-4"),
    "noisy_or": ("Srinivas — A Generalization of the Noisy-Or Model", "https://doi.org/10.1016/B978-1-4832-1451-1.50030-5"),
    "ratcliff": ("Ratcliff and Metzener — Pattern Matching: The Gestalt Approach", "https://jacobfilipp.com/DrDobbs/articles/DDJ/1988/8807/8807c/8807c.htm"),
    "kleene": ("Kleene — Representation of Events in Nerve Nets and Finite Automata", "https://doi.org/10.1515/9781400882618-002"),
    "rfc9562": ("IETF RFC 9562 — Universally Unique IDentifiers", "https://doi.org/10.17487/RFC9562"),
}

SOURCE_DOCS = [
    r"argus_main\docs\ArgusWatch_Project_Professional_Documentation.docx",
    r"cyber_threat_monitor\docs\Cyber_Threat_Monitor_Report_updated.docx",
    r"osint and breach codebase\docs\OSINT module v1.docx",
    r"osint and breach codebase\docs\Breach Module v1.docx",
    r"pii_link\docs\pii.docx",
]


# Matrices are deliberately conservative. A check means the capability is
# implemented in this repository or explicitly described by the cited primary
# source. A cross means it was not evidenced as a native capability during the
# 10 August 2026 review; it does not rule out custom extensions or partners.
CORE_COMPARISON = (
    ["Feature", "Argus Unified", "OpenCTI", "MISP", "ThreatConnect", "Recorded Future"],
    [
        ["Self-hosted open-source deployment", "✓", "✓", "✓", "✗", "✗"],
        ["Structured CTI objects and relationships", "✓", "✓", "✓", "✓", "✓"],
        ["Feed/import connector framework", "✓", "✓", "✓", "✓", "✓"],
        ["STIX 2.1-compatible export", "✓", "✓", "✓", "✓", "✗"],
        ["Customer/victim asset inventory", "✓", "✓", "✗", "✓", "✓"],
        ["Automated asset-to-indicator correlation", "✓", "✓", "✗", "✗", "✓"],
        ["Native attack-surface/recon collection", "✓", "✗", "✗", "✗", "✓"],
        ["Native breach-record search and graph", "✓", "✗", "✗", "✗", "✗"],
        ["Native Telegram collection and analysis", "✓", "✗", "✗", "✗", "✗"],
        ["Case/task or investigation workflow", "✓", "✓", "✓", "✓", "✗"],
        ["Finding SLA and remediation state", "✓", "✗", "✗", "✗", "✗"],
        ["Published deterministic exposure formula", "✓", "✗", "✗", "✗", "✗"],
        ["Built-in commercial intelligence corpus", "✗", "✗", "✗", "✗", "✓"],
        ["API/SIEM-oriented integration", "✓", "✓", "✓", "✓", "✓"],
        ["RBAC and auditable activity controls", "✓", "✓", "✓", "✓", "✗"],
        ["Dashboards or analyst reporting", "✓", "✓", "✓", "✓", "✓"],
    ],
)

OSINT_COMPARISON = (
    ["Feature", "Argus Unified", "SpiderFoot", "Maltego", "Sherlock", "Maigret"],
    [
        ["Self-hosted open-source core", "✓", "✓", "✗", "✓", "✓"],
        ["Username enumeration", "✓", "✓", "✓", "✓", "✓"],
        ["Name/email/phone/person inputs", "✓", "✓", "✓", "✗", "✗"],
        ["Multi-source module/transform ecosystem", "✓", "✓", "✓", "✗", "✗"],
        ["Rich subject-profile disambiguation", "✓", "✗", "✗", "✗", "✗"],
        ["Evidence-chain confidence tiers", "✓", "✗", "✗", "✗", "✗"],
        ["DOM-skeleton account verification", "✓", "✗", "✗", "✗", "✗"],
        ["Structured profile metadata extraction", "✓", "✓", "✓", "✗", "✓"],
        ["Recursive pivot discovery", "✓", "✓", "✓", "✗", "✓"],
        ["WAF/challenge-aware checks", "✓", "✗", "✗", "✗", "✓"],
        ["Known-positive/negative calibration", "✓", "✗", "✗", "✗", "✗"],
        ["Pause/resume/cancel background jobs", "✓", "✗", "✗", "✗", "✗"],
        ["Visual relationship graph", "✗", "✓", "✓", "✗", "✓"],
        ["Real-time graph collaboration", "✗", "✗", "✓", "✗", "✗"],
        ["Local CSV/intelbase lookup", "✓", "✗", "✓", "✗", "✗"],
        ["Structured report/export formats", "✓", "✓", "✓", "✓", "✓"],
    ],
)

BREACH_COMPARISON = (
    ["Feature", "Argus Unified", "HIBP", "Intelligence X", "LeakCheck", "DeHashed"],
    [
        ["Self-hosted local-corpus search", "✓", "✗", "✗", "✗", "✗"],
        ["Vendor-hosted breach corpus", "✗", "✓", "✓", "✓", "✓"],
        ["Email lookup", "✓", "✓", "✓", "✓", "✓"],
        ["Username lookup", "✓", "✓", "✓", "✓", "✓"],
        ["Phone lookup", "✓", "✓", "✓", "✓", "✓"],
        ["IP-address lookup", "✓", "✗", "✓", "✗", "✓"],
        ["Domain exposure/search", "✓", "✓", "✓", "✓", "✗"],
        ["Name/address pivot", "✓", "✗", "✗", "✗", "✓"],
        ["Combined multi-field query", "✓", "✗", "✗", "✗", "✓"],
        ["Full raw record fields", "✓", "✗", "✓", "✓", "✓"],
        ["Dataset/breach provenance", "✓", "✓", "✓", "✓", "✓"],
        ["Paginated programmatic API", "✓", "✗", "✓", "✓", "✓"],
        ["Record-to-PII relationship graph", "✓", "✗", "✗", "✗", "✗"],
        ["Common-value graph expansion guard", "✓", "✗", "✗", "✗", "✗"],
        ["Stable local node/edge identifiers", "✓", "✗", "✗", "✗", "✗"],
        ["Redis cache with uncached fallback", "✓", "✗", "✗", "✗", "✗"],
        ["Integrated customer finding/exposure", "✓", "✗", "✗", "✗", "✗"],
        ["Hosted monitoring/notifications", "✗", "✓", "✓", "✗", "✓"],
    ],
)

TELEGRAM_COMPARISON = (
    ["Feature", "Argus Unified", "Telethon", "Maltego Monitor", "Social Links", "OpenCTI"],
    [
        ["Self-hosted open-source component", "✓", "✓", "✗", "✗", "✓"],
        ["Authorized Telegram account/API collection", "✓", "✓", "✓", "✓", "✗"],
        ["Explicit channel/group monitoring", "✓", "✓", "✓", "✓", "✗"],
        ["Historical message retrieval", "✓", "✓", "✓", "✓", "✗"],
        ["Incoming-message event handling", "✗", "✓", "✓", "✗", "✗"],
        ["Bounded JSON/ZIP offline import", "✓", "✗", "✗", "✗", "✗"],
        ["Text, links and attachment extraction", "✓", "✓", "✓", "✓", "✗"],
        ["Obfuscation/leetspeak normalization", "✓", "✗", "✗", "✗", "✗"],
        ["Deterministic cybercrime categories", "✓", "✗", "✗", "✗", "✗"],
        ["Published risk/confidence formula", "✓", "✗", "✗", "✗", "✗"],
        ["Typo-tolerant local message search", "✓", "✗", "✗", "✗", "✗"],
        ["Alert/case monitoring interface", "✓", "✗", "✓", "✓", "✓"],
        ["Channel/reference relationship graph", "✓", "✗", "✗", "✓", "✓"],
        ["Customer asset-to-IOC findings", "✓", "✗", "✗", "✗", "✓"],
        ["STIX 2.1 export after ingestion", "✓", "✗", "✗", "✗", "✓"],
        ["General-purpose Telegram client API", "✗", "✓", "✗", "✗", "✗"],
    ],
)

PII_COMPARISON = (
    ["Feature", "Argus Unified", "Presidio", "Google SDP", "Amazon Macie"],
    [
        ["Self-hosted text detection", "✓", "✓", "✗", "✗"],
        ["Built-in PII/financial identifiers", "✓", "✓", "✓", "✓"],
        ["Credential/secret detection", "✓", "✗", "✓", "✓"],
        ["Regex/context/custom detectors", "✓", "✓", "✓", "✓"],
        ["NER or machine-learning detection", "✗", "✓", "✓", "✓"],
        ["Detection confidence/likelihood", "✓", "✓", "✓", "✗"],
        ["Threat IOC families in same scan", "✓", "✗", "✗", "✗"],
        ["Cloud data-store discovery", "✗", "✗", "✓", "✓"],
        ["Image PII inspection/redaction", "✗", "✓", "✓", "✗"],
        ["Anonymize/de-identify operations", "✗", "✓", "✓", "✗"],
        ["Customer asset attribution", "✓", "✗", "✗", "✗"],
        ["Breach relationship-graph integration", "✓", "✗", "✗", "✗"],
        ["STIX 2.1 CTI export", "✓", "✗", "✗", "✗"],
        ["Fully managed cloud service", "✗", "✗", "✓", "✓"],
    ],
)

IOC_STIX_COMPARISON = (
    ["Feature", "Argus Unified", "iocextract", "OpenCTI", "MISP"],
    [
        ["Self-hosted open-source component", "✓", "✓", "✓", "✓"],
        ["IOC extraction from free text", "✓", "✓", "✓", "✓"],
        ["Defanged IOC decoding/refanging", "✗", "✓", "✗", "✗"],
        ["Credential/token/PII pattern families", "✓", "✗", "✗", "✗"],
        ["Normalization and deduplication", "✓", "✗", "✓", "✓"],
        ["Per-indicator confidence", "✓", "✗", "✓", "✓"],
        ["False-positive/noise guard lists", "✓", "✗", "✗", "✓"],
        ["Customer asset/CIDR/domain matching", "✓", "✗", "✓", "✗"],
        ["CVE-to-product/version correlation", "✓", "✗", "✗", "✗"],
        ["STIX 2.1 import/export platform", "✓", "✗", "✓", "✓"],
        ["Deterministic UUIDv5 output IDs", "✓", "✗", "✗", "✗"],
        ["Knowledge/event relationship graph", "✓", "✗", "✓", "✓"],
        ["Breach-record graph integration", "✓", "✗", "✗", "✗"],
        ["Case or workflow handling", "✓", "✗", "✓", "✓"],
    ],
)


def _font(name: str = "Calibri", size: float = 11, bold: bool = False, color: str = BLACK):
    return {"name": name, "size": Pt(size), "bold": bold, "color": RGBColor.from_string(color)}


def _set_run(run, *, name="Calibri", size=11, bold=False, italic=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[idx])
            _cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def _hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    link.append(run)
    paragraph._p.append(link)


def _field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def _configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 24, BLACK, 0, 10),
        ("Subtitle", 12, MID_GRAY, 0, 10),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [s.name for s in styles]:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0
    p_pr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)


def _configure_page(doc: Document, module_name: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_table_geometry(table, [4680, 4680])
    for cell in table.rows[0].cells:
        _cell_margins(cell, 0, 0, 0, 0)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ARGUS UNIFIED | IMPLEMENTATION REPORT")
    _set_run(r, size=8, bold=True, color=DARK_BLUE)
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(module_name.upper())
    _set_run(r, size=8, bold=True, color=MID_GRAY)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_table_geometry(table, [7000, 2360])
    for cell in table.rows[0].cells:
        _cell_margins(cell, 0, 0, 0, 0)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Verified implementation · {DATE}")
    _set_run(r, size=8, color=MID_GRAY)
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Page ")
    _set_run(r, size=8, color=MID_GRAY)
    _field(p, "PAGE")


def new_document(module_name: str) -> Document:
    doc = Document()
    _configure_styles(doc)
    _configure_page(doc, module_name)
    props = doc.core_properties
    props.title = module_name
    props.subject = "Argus Unified verified implementation documentation"
    props.author = "Argus Unified Engineering"
    props.keywords = "threat intelligence, OSINT, breach, Telegram, IOC, STIX"
    return doc


def masthead(doc: Document, title: str, subtitle: str, status: str = "IMPLEMENTED / VERIFIED"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("ARGUS UNIFIED")
    _set_run(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    rows = [
        ("Document status", status),
        ("As of", DATE),
        ("Basis", "Consolidated source code, automated tests, build outputs, and browser verification"),
        ("Scope rule", "Implemented behavior only; no roadmap or unimplemented product claims"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [2050, 7310])
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = key
        table.cell(idx, 1).text = value
        _shade(table.cell(idx, 0), LIGHT_BLUE)
        for run in table.cell(idx, 0).paragraphs[0].runs:
            _set_run(run, size=9, bold=True, color=DARK_BLUE)
        for run in table.cell(idx, 1).paragraphs[0].runs:
            _set_run(run, size=9)
    doc.add_paragraph()


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        _set_run(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        _set_run(r, italic=italic)
    else:
        r = p.add_run(text)
        _set_run(r, italic=italic)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code(doc: Document, lines: str):
    for line in lines.strip().splitlines():
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.keep_together = True
        p.add_run(line)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    _set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        _shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            _set_run(run, size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            if value in ("✓", "✗"):
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cells[idx].paragraphs[0].runs:
                    _set_run(run, name="Segoe UI Symbol", size=font_size + 1, bold=True)
            else:
                for run in cells[idx].paragraphs[0].runs:
                    _set_run(run, size=font_size)
        _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_sources(doc: Document, keys: list[str]):
    doc.add_heading("Reviewed external sources", level=2)
    add_para(doc, "Competitor capabilities were marked only when explicitly evidenced by the reviewed official source. Accessed 10 August 2026.")
    for key in keys:
        label, url = SOURCES[key]
        p = doc.add_paragraph(style="List Bullet")
        _hyperlink(p, label, url)


def add_lineage(doc: Document, used: list[str]):
    doc.add_heading("Source lineage and reuse", level=2)
    add_para(doc, "The consolidated report reuses verified implementation descriptions, workflows, and terminology from the existing project documents, reconciled against the final code. Obsolete database stacks, duplicate services, future work, and unverified readiness claims were excluded.")
    add_bullets(doc, used)


def add_comparison(doc: Document, headers: list[str], rows: list[list[str]], *, title="Market feature comparison"):
    doc.add_heading(title, level=2)
    add_para(doc, "Evidence rule (reviewed 10 August 2026): ✓ = native capability implemented in Argus or explicitly documented by the cited primary source; ✗ = not evidenced as a native capability in the reviewed source. A cross does not mean a custom integration, partner feed, or extension is impossible. Product editions and licensing can change.")
    feature_width = 3360
    remaining = CONTENT_DXA - feature_width
    col = remaining // (len(headers) - 1)
    widths = [feature_width] + [col] * (len(headers) - 2) + [remaining - col * (len(headers) - 2)]
    add_table(doc, headers, rows, widths, font_size=7.2 if len(headers) >= 6 else 7.7)


def add_research_basis(doc: Document, keys: list[str], caveat: str, *, level: int = 2):
    doc.add_heading("Research and standards basis", level=level)
    add_para(doc, caveat)
    for key in keys:
        label, url = SOURCES[key]
        p = doc.add_paragraph(style="List Bullet")
        _hyperlink(p, label, url)


def add_core_math(doc: Document, *, level: int = 1):
    sub = min(level + 1, 3)
    doc.add_heading("Flow-aligned mathematical algorithm and specification", level=level)
    add_para(doc, "This specification follows the core flowchart from normalized evidence through threat pressure, customer correlation, exposure scoring, and bounded operational output. Every coefficient below is present in the repository; no coefficient is presented as learned or statistically calibrated.")
    add_table(doc, ["Symbol", "Range / type", "Implemented meaning"], [
        ["s_raw", "real, normally 0..100", "Raw evidence or source score"],
        ["c_src", "0..1", "Configured feed confidence; unknown source defaults to 0.5"],
        ["a, h", "days; h > 0", "Evidence age and decay half-life; default h=14, pressure h=7"],
        ["D1..D5", "each 0..100", "Direct exposure, active exploitation, actor intent, attack surface, asset criticality"],
        ["n", "positive integer", "Observation count in a threat category/family group"],
        ["P_j", "0..10", "Decayed activity of threat group j"],
    ], [1200, 2050, 6110], font_size=8.2)

    doc.add_heading("Freshness and source reliability", level=sub)
    add_code(doc, "lambda = ln(2) / h\nw_time(a,h) = exp(-lambda * max(a,0))\ns_norm = s_raw * c_src * w_time(a,h)\n\nInvariant: w_time(0,h)=1 and w_time(h,h)=0.5")
    add_para(doc, "The half-life parameter makes the decay interpretable and monotone: older evidence never gains weight, and every h days the temporal contribution halves. The implementation does not delete old evidence; it reduces its contribution to the current score.")

    doc.add_heading("Threat activity and sector pressure", level=sub)
    add_code(doc, "A(n) = min(10, 3.3 * log10(max(n,1)))\nb(r) = 1.3 if age_hours < 1; 1.1 if age_hours < 6; otherwise 1\nP_j = min(10, A(n_j) * b(r_j) * exp(-ln(2) * age_days_j / 7))\nP_sector = min(10, 0.60 * max_j(P_j) + 0.40 * mean_j(P_j))")
    add_para(doc, "The logarithm compresses burst volume, the recency multiplier distinguishes current activity, and the max/mean blend preserves both the dominant threat and the breadth of targeting. Groups with fewer than two observations are not persisted as pressure groups.")

    doc.add_heading("Five-dimension exposure score", level=sub)
    add_code(doc, "E_signal = 0.50*D1 + 0.30*D2 + 0.20*D3\nE_floor  = 0.20*D4\nE_base   = max(E_signal, E_floor)\nM_impact = 0.75 + 0.00125*D4 + 0.00125*D5\nRisk     = min(100, E_base * M_impact)")
    add_para(doc, "For D4,D5 in [0,100], M_impact is bounded in [0.75,1.00]. Risk is bounded in [0,100], is monotone non-decreasing in every dimension, and cannot be zero solely because direct threat evidence is absent when attack-surface exposure exists. Operational labels are LOW <40, MEDIUM 40-59, HIGH 60-79, and CRITICAL 80-100.")

    doc.add_heading("Executable algorithm", level=sub)
    add_code(doc, "CORE-SCORE(evidence, customer):\n  1. validate scope, tenant and provenance\n  2. normalize and deduplicate evidence\n  3. for each item: apply source confidence and half-life decay\n  4. group current observations; compute P_j and P_sector\n  5. correlate exact customer assets and compute D1..D5\n  6. compute E_base, M_impact and bounded Risk\n  7. persist finding, SLA/remediation state and score history\n  8. emit reviewed alerts, reports, STIX or SIEM output")
    add_table(doc, ["Property", "Result"], [
        ["Determinism", "Same stored evidence, customer assets, configuration, and time reference produce the same score"],
        ["Boundedness", "Pressure <=10 and final Risk <=100"],
        ["Dominant runtime", "Linear in evidence processed plus datastore query/index costs; scoring arithmetic is O(N)"],
        ["Calibration boundary", "Weights and thresholds are engineering policy constants, not fitted probabilities"],
    ], [2200, 7160], font_size=8.2)
    add_research_basis(doc, ["nist_risk", "temporal_ir"], "NIST SP 800-30 supports explicit treatment of threat sources, vulnerabilities, likelihood and impact, while temporal-information-retrieval research supports time-aware weighting. Neither source specifies Argus's exact coefficients; the formula is an auditable implementation policy that still requires organization-specific validation and sensitivity analysis.", level=sub)


def add_osint_math(doc: Document, *, level: int = 1):
    sub = min(level + 1, 3)
    doc.add_heading("Flow-aligned mathematical algorithm and specification", level=level)
    add_para(doc, "The OSINT mathematics maps the flowchart's subject profile, direct checks, response verification, evidence chain, confidence tier, correlation, and analyst-review stages. Scores rank evidence; they do not prove identity.")
    doc.add_heading("Subject evidence and profile-strength priors", level=sub)
    add_code(doc, "S_subject = min(100, sum_j w_j * I_j)\nI_j = 1 when the normalized candidate contains the supplied attribute, else 0")
    add_table(doc, ["Matched signal", "Implemented weight"], [
        ["Email; phone", "+30 each (first match of each type)"],
        ["Workplace; company", "+25; +20 (company counted once)"],
        ["Education; occupation; industry", "+20; +15; +10"],
        ["City; state; non-obvious country", "+20; +10; +5"],
        ["Full display name; known-profile cross-link", "+10; +20"],
    ], [4200, 5160], font_size=8.2)
    add_code(doc, "S_profile = min(100, 35*email + 30*phone + 15*photo + 20*username\n                         + 15*workplace + 10*education + 10*city + 5*full_name)\nSTRONG >=60; MODERATE 25..59; WEAK <25")
    add_para(doc, "All variables in S_profile are binary presence indicators. The score measures expected disambiguation power of the supplied input, not the probability that any result is correct.")

    doc.add_heading("DOM soft-404 discrimination", level=sub)
    add_code(doc, "For tag t, let f_A(t), f_B(t) be counts in the first 10 KB.\nJ_w(A,B) = sum_t min(f_A(t),f_B(t)) / sum_t max(f_A(t),f_B(t))\nDOM differs when J_w < 0.90")
    add_para(doc, "This is a weighted multiset form of Jaccard similarity. A candidate page that closely resembles the calibrated nonexistent-account page is treated as soft-404 evidence rather than a confirmed profile.")

    doc.add_heading("Rank prior and three-way decision", level=sub)
    add_code(doc, "R_on(k)  = max(0, floor(35*(1-k/12)))\nR_off(k) = max(0, floor(25*(1-k/12)))\nFinal rank = min(100, rank prior + documented attribute boosts)\n\nDecision = CONFIRMED/HIGH/MEDIUM/AMBIGUOUS/UNVERIFIED/NOT_FOUND\nThresholds = 95/75/50/30/1/0")
    add_para(doc, "The evidence-chain gate separately records DEFINITIVE, HIGH, POSSIBLE, or UNLIKELY support. This preserves a link / possible-link / non-link style decision rather than forcing every response into a binary identity claim.")

    doc.add_heading("Executable algorithm", level=sub)
    add_code(doc, "OSINT-INVESTIGATE(profile):\n  1. validate scope; compute permutations and S_profile\n  2. run bounded direct-site, SERP and local-intelbase checks\n  3. classify status, redirects, markers and WAF/challenge evidence\n  4. compare candidate DOM with calibrated negative baseline\n  5. extract metadata and compute S_subject plus evidence tier\n  6. assign confidence class; keep negative and ambiguous evidence\n  7. cluster cross-platform candidates and return bounded pivots\n  8. require analyst review before attribution")
    add_table(doc, ["Property", "Result"], [
        ["Score bounds", "All public scores are clipped to 0..100"],
        ["DOM cost", "O(T_A + T_B) over extracted tag counts"],
        ["Network cost", "Dominated by bounded external requests, retries and configured concurrency"],
        ["Statistical boundary", "Weights are deterministic heuristics, not Fellegi-Sunter log-likelihood ratios learned from labeled pairs"],
    ], [2200, 7160], font_size=8.2)
    add_research_basis(doc, ["fellegi_sunter", "jaccard"], "Fellegi-Sunter provides the classical theoretical basis for link, possible-link, and non-link decisions from comparison evidence. Jaccard provides the similarity basis adapted here to DOM tag-frequency multisets. Argus borrows these structures but does not claim a trained probabilistic record-linkage model.", level=sub)


def add_breach_math(doc: Document, *, level: int = 1):
    sub = min(level + 1, 3)
    doc.add_heading("Flow-aligned mathematical algorithm and specification", level=level)
    add_para(doc, "The breach algorithm follows the flowchart's validated seed, controlled field catalog, cache decision, parameterized search, provenance-preserving records, entity extraction, bounded graph traversal, and capped response.")
    doc.add_heading("Canonicalization and stable identifiers", level=sub)
    add_code(doc, "email_norm(x) = lowercase(trim(x))\nphone_norm(x) = canonical ten-digit representation after prefix/trunk handling\nrow_id    = SHA256(source_file || '|' || email_norm || '|' || phone_norm)\nentity_id = SHA256(UPPER(type) || ':' || lowercase(trim(value)))\nedge_id   = SHA256(row_id || '|' || entity_id)")
    add_para(doc, "These hashes make repeated rows and entities stable within the same canonicalization policy. They are identifiers, not proof that two people are identical and not a substitute for encrypting sensitive values.")

    doc.add_heading("Bipartite graph and bounded BFS", level=sub)
    add_code(doc, "G = (R union E, L)\nR = breach-record nodes; E = normalized EMAIL/PHONE/NAME/ADDRESS nodes\n(r,e) in L iff record r contains entity e\n\nHard bounds: |R| <= 500, |E| <= 300, |R union E| <= 800\nCommon-value gate: degree(e) > 200 => retain 10 samples and do not expand\nExpansion depth: seed records -> entities -> first-degree records -> second-degree entities")
    add_code(doc, "BREACH-GRAPH(seed):\n  1. normalize and resolve at most 100 seed records\n  2. create stable record/entity nodes and enqueue unseen entities\n  3. dequeue entities in batches of at most 10\n  4. query 201 rows to detect degree >200\n  5. if common: mark warning, keep 10 samples, stop that branch\n  6. else add deduplicated rows and one further entity layer\n  7. stop at 500 records, 300 entities, timeout, or empty queue\n  8. return provenance, pagination, capped and cap_reason metadata")
    add_table(doc, ["Property", "Result"], [
        ["Traversal", "Breadth-first over a bipartite graph with explicit visited sets"],
        ["In-memory complexity", "O(|R|+|E|+|L|), absolutely bounded by configured node caps"],
        ["External query cost", "Bounded by unique expanded entities, batch size, timeout and datastore indexes"],
        ["Precision guard", "High-degree values are warned and sampled because common names/addresses create noisy joins"],
        ["Interpretation boundary", "An edge shows shared data, not identity, ownership, compromise, or causation"],
    ], [2200, 7160], font_size=8.2)
    add_research_basis(doc, ["moore_bfs", "meta_blocking", "fips180"], "Moore's breadth-first search supplies the traversal basis; entity-resolution meta-blocking research supports restricting noisy comparison graphs for scalability; FIPS 180-4 specifies SHA-256. Argus's thresholds (200, 10, 500, 300 and depth two) are implementation safety bounds, not values established as optimal by those publications.", level=sub)


def add_telegram_math(doc: Document, *, level: int = 1):
    sub = min(level + 1, 3)
    doc.add_heading("Flow-aligned mathematical algorithm and specification", level=level)
    add_para(doc, "The Telegram algorithm follows authorization/import gates, normalization, rule categories, saturating confidence, bounded risk aggregation, flagging, persistence, customer correlation, search, graphing, and STIX export.")
    doc.add_heading("Category saturation and aggregate risk", level=sub)
    add_code(doc, "r_k = sum_i w_i * I(rule_i matches raw text or normalized text)\nc_k = 1 - exp(-r_k/4)\nA   = sum_k alpha_k * c_k\np_0 = 1 - exp(-A)")
    add_table(doc, ["Category k", "alpha_k"], [
        ["Ransomware", "1.00"], ["Data breach", "0.90"], ["Initial access", "0.90"],
        ["Carding", "0.85"], ["Phishing", "0.75"], ["Exploit", "0.80"], ["Broker contact", "0.35"],
    ], [6000, 3360], font_size=8.2)
    add_code(doc, "if external channel: p = 1 - (1-p_0)*0.90\nif risky attachment: p = 1 - (1-p)*0.88\nrisk = clip(round(100*p), 0, 100)\nflagged = (risk >= 55) OR (c_ransomware >= 0.65)")
    add_para(doc, "The exponential transforms are monotone and bounded, so repeated synonymous rules have diminishing marginal effect. The complement updates have a noisy-OR-style form. The resulting p is a risk-like score, not a calibrated real-world probability of criminality or compromise.")

    doc.add_heading("Typo-tolerant token similarity", level=sub)
    add_code(doc, "sim(q,t) = 2*M / (|q|+|t|)\nM = total characters in recursively matched common subsequences\nbest(q,text) = max over bounded query/candidate token pairs sim(q_i,t_j)")
    add_para(doc, "The implementation uses Python SequenceMatcher, discards tokens shorter than four characters, caps candidate tokens at 450, and skips pairs whose lengths differ by more than 50 percent. Search fetches at most 2,000 database candidates before ranking.")

    doc.add_heading("Executable algorithm", level=sub)
    add_code(doc, "TELEGRAM-ANALYZE(message):\n  1. validate authorized session/channel or bounded import\n  2. normalize case, leetspeak, separators and whitespace\n  3. evaluate deterministic rules in raw and normalized text\n  4. compute r_k, c_k, A and bounded risk\n  5. apply channel/attachment complement boosts; evaluate flag gate\n  6. extract links, mentions, attachments and IOCs\n  7. persist evidence and correlate exact customer assets\n  8. expose alerts/search/graph and STIX after analyst review")
    add_table(doc, ["Property", "Result"], [
        ["Bounds", "0 <= c_k < 1 and 0 <= risk <= 100"],
        ["Rule evaluation", "Deterministic for the same text, media name and rule set"],
        ["Fuzzy-search bound", "At most 450 candidate tokens per message and 2,000 fetched message rows"],
        ["Validation boundary", "Weights and flag thresholds require labeled-corpus precision/recall calibration before probabilistic interpretation"],
    ], [2200, 7160], font_size=8.2)
    add_research_basis(doc, ["noisy_or", "ratcliff"], "Noisy-OR literature supports bounded aggregation of multiple partially independent causes, and Ratcliff-Metzener describes the Gestalt string-similarity family used by SequenceMatcher. Argus uses these mathematical shapes deterministically; independence and probability calibration have not been established on a labeled Telegram corpus.", level=sub)


def add_pii_math(doc: Document, *, level: int = 1):
    sub = min(level + 1, 3)
    doc.add_heading("Flow-aligned mathematical algorithm and specification", level=level)
    add_para(doc, "The PII/IOC/STIX algorithm follows ordered pattern recognition, noise guards, canonicalization, deduplication, confidence adjustment, cautious customer attribution, and standards-based object mapping.")
    doc.add_heading("Ordered detection and confidence", level=sub)
    add_code(doc, "For ordered pattern p_i and line l:\nI_i(l) = 1 if regex p_i matches l and all noise guards pass, else 0\nK(match) = (category, ioc_type, exact_value)\nkeep the first unique K under specific-to-generic scan order\nconfidence = min(1, base_confidence + 0.10*I(customer_domain occurs on same line))")
    add_para(doc, "The engine evaluates 92 expressions in 17 ordered families. Specific credential, token, session and exfiltration patterns precede generic URL, domain, email, IP and hash patterns. Private/localhost IP guards and removed high-noise expressions reduce false positives.")

    doc.add_heading("Customer attribution predicates", level=sub)
    add_code(doc, "A(e,c) = IP_EXACT_OR_CIDR(e,c) OR DOMAIN_EXACT_OR_SUBDOMAIN(e,c)\n         OR EMAIL_DOMAIN_OR_EXECUTIVE(e,c) OR CVE_PRODUCT_VERSION(e,c)\n         OR BOUNDED_BRAND_KEYWORD_EVIDENCE(e,c)\n\nDirect finding only when an implemented attribution predicate supplies match proof.\nUnknown product version => probable exposure at lower confidence.\nNo attributable predicate => global signal, not a customer finding.")
    add_para(doc, "Domain matching is boundary-aware; raw substring containment is not treated as exact ownership. IP ranges use canonical address membership. These are conservative logical predicates rather than a learned entity-resolution model.")

    doc.add_heading("STIX confidence and deterministic names", level=sub)
    add_code(doc, "C_stix = clip(round(100*confidence), 0, 100)\nUUIDv5(namespace,name) = RFC9562_Format(first128(SHA1(namespace || canonical_name)))\nBundle UUIDs are UUIDv4; Telegram-derived object IDs use UUIDv5.\nDuplicate exported objects are removed by normalized (type,value) identity.")
    add_para(doc, "Standard IP, domain, URL, email and hash values become STIX Indicator patterns; CVEs become Vulnerability objects in the Telegram export path; unsupported types use the documented open 'argus' pattern type instead of pretending to be valid STIX pattern syntax.")

    doc.add_heading("Executable algorithm", level=sub)
    add_code(doc, "PII-IOC-STIX(text, customer):\n  1. scan each line with specific-to-generic compiled patterns\n  2. reject short, private, localhost and duplicate candidates\n  3. normalize values; attach line context and bounded confidence\n  4. evaluate exact/boundary/CIDR/CVE customer predicates\n  5. create direct, probable or global evidence without false attribution\n  6. map supported values to STIX 2.1 objects and escape patterns\n  7. assign bounded confidence and stable IDs where implemented\n  8. deduplicate bundle and retain evidence for analyst review")
    add_table(doc, ["Property", "Result"], [
        ["Detection cost", "O(L * P * match_cost) for L lines and P compiled expressions; P is fixed at 92"],
        ["Deduplication", "Expected O(M) with a hash set for M candidate matches"],
        ["Interoperability", "STIX structure and pattern vocabulary follow OASIS STIX 2.1"],
        ["Validation boundary", "Regex-shaped PII and secrets are candidates; provider checksums/contextual validation may still be required"],
    ], [2200, 7160], font_size=8.2)
    add_research_basis(doc, ["kleene", "stix", "rfc9562", "fips180"], "Kleene's formal-language work supplies the theoretical basis for regular-expression recognition; OASIS specifies STIX 2.1; RFC 9562 specifies UUIDv5; FIPS 180-4 specifies SHA-1/SHA-256 primitives. These standards validate representation mechanics, not the semantic truth of a regex match or customer attribution.", level=sub)


def add_unified_math(doc: Document, *, level: int = 1):
    sub = min(level + 1, 3)
    doc.add_heading("Unified end-to-end mathematical algorithm", level=level)
    add_para(doc, "The unified model composes the five module algorithms without averaging incompatible scores. OSINT identity evidence, breach relationships, Telegram risk, IOC candidates and core exposure remain typed values until explicit correlation predicates create a customer finding.")
    add_code(doc, "O = O_core union O_osint union O_breach union O_telegram union O_pii_ioc\nE = DEDUP(NORMALIZE(O), provenance-preserving typed key)\nF_c = { e in E : A(e,c)=1 and policy_gate(e,c)=ALLOW }\nG_c = { e in E : A(e,c)=0 }  # global/sector evidence, never forced attribution\nRisk_c = CORE_D1_D5(F_c, threat_pressure(G_c), assets_c)")
    add_code(doc, "ARGUS-UNIFIED(input, customer):\n  1. enforce authorization, tenant, source and rate-limit gates\n  2. dispatch typed input to OSINT, breach, Telegram and IOC modules\n  3. preserve each module's raw evidence, confidence and provenance\n  4. normalize/deduplicate; do not numerically merge unlike scores\n  5. evaluate exact customer-asset attribution predicates\n  6. route direct evidence to findings and unmatched evidence to pressure\n  7. calculate D1..D5 exposure and remediation priority\n  8. require analyst review for identity, attribution and consequential action\n  9. export reviewed intelligence to STIX/SIEM/reports/alerts\n 10. retain audit trail, score history and operational bounds")
    add_table(doc, ["System invariant", "Meaning"], [
        ["Type preservation", "An OSINT confidence class is not treated as a Telegram probability or breach graph weight"],
        ["No forced attribution", "Unmatched indicators contribute only to global/sector pressure"],
        ["Bounded outputs", "Module caps and the 0..100 core score prevent unbounded operational values"],
        ["Human decision gate", "Automated evidence ranking does not establish identity, intent, ownership or legal conclusion"],
        ["Reproducibility", "Deterministic modules reproduce results for identical inputs, configuration, data and time reference"],
    ], [2200, 7160], font_size=8.2)
    add_research_basis(doc, ["nist_risk", "fellegi_sunter", "meta_blocking", "noisy_or", "stix"], "The complete pipeline combines established structures from risk assessment, record linkage, bounded entity-resolution graphs, noisy-OR-style evidence aggregation and standardized CTI representation. The integration is engineering work specific to Argus; publication-quality empirical validation would additionally require labeled datasets, calibration curves, error analysis, ablation studies and external replication.", level=sub)


def add_verification(doc: Document, module_specific: list[str]):
    doc.add_heading("Verification evidence", level=2)
    items = [
        "All 193 Python source files parse successfully.",
        "Python unit/stateless suite: 242 tests passed in the isolated Python 3.12 environment.",
        "Application smoke subset: 9 tests passed, including boot, CORS, validation, and authentication flows.",
        "Frontend ESLint and optimized Vite build completed successfully; npm audit reported zero known vulnerabilities.",
        "pip-audit reported zero known vulnerabilities for backend/requirements.txt.",
    ] + module_specific
    add_bullets(doc, items)


def add_boundaries(doc: Document, items: list[str]):
    heading = doc.add_heading("Current operational boundaries", level=2)
    heading.paragraph_format.keep_with_next = True
    for idx, item in enumerate(items):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        if idx < len(items) - 1:
            p.paragraph_format.keep_with_next = True


def add_flow(doc: Document, flow_name: str, caption: str):
    """Add a flowchart with a PNG compatibility preview.

    ``save_with_embedded_svg`` upgrades the drawing to a native SVG part after
    python-docx saves the package. Modern Word uses the SVG; older readers keep
    the PNG fallback instead of displaying a broken placeholder.
    """
    path = ASSETS / f"{flow_name}.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.15))
    shape._inline.docPr.set("name", flow_name)
    shape._inline.docPr.set("descr", caption)
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        _set_run(run, size=8, italic=True, color=MID_GRAY)
    # A detailed architecture chart intentionally owns its page. Starting the
    # following section on a fresh page prevents orphaned table headers below
    # the figure and keeps the page-2-inspired presentation intact.
    p.add_run().add_break(WD_BREAK.PAGE)


def _embed_svg_fallback(docx_path: Path, svg_path: Path, flow_name: str):
    """Attach an SVG to the existing PNG drawing using Word's SVG extension.

    The standard DrawingML ``a:blip`` continues to reference the PNG fallback.
    ``asvg:svgBlip`` points at the vector part, which keeps text and connectors
    sharp at any zoom in Microsoft Word while preserving compatibility.
    """
    relationship_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    wordprocessing_drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    svg_ns = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
    content_type_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    image_rel_type = f"{office_rel_ns}/image"

    with zipfile.ZipFile(docx_path, "r") as source:
        package = {item.filename: source.read(item.filename) for item in source.infolist()}

    rel_path = "word/_rels/document.xml.rels"
    document_path = "word/document.xml"
    rel_root = etree.fromstring(package[rel_path])
    document_root = etree.fromstring(package[document_path])
    named_blips = document_root.xpath(
        ".//wp:inline[wp:docPr[@name=$flow_name]]//a:blip",
        namespaces={"wp": wordprocessing_drawing_ns, "a": drawing_ns},
        flow_name=flow_name,
    )
    if len(named_blips) != 1:
        raise RuntimeError(f"Expected one named flowchart drawing for {flow_name}, found {len(named_blips)}")
    blip = named_blips[0]
    fallback_id = blip.get(f"{{{office_rel_ns}}}embed")
    fallback_candidates = [
        rel
        for rel in rel_root.findall(f"{{{relationship_ns}}}Relationship")
        if rel.get("Type") == image_rel_type
        and rel.get("Id") == fallback_id
        and rel.get("Target", "").lower().endswith(".png")
    ]
    if len(fallback_candidates) != 1:
        raise RuntimeError(
            f"Expected one PNG flowchart relationship for {flow_name}, found {len(fallback_candidates)}"
        )
    fallback_rel = fallback_candidates[0]

    used_ids = {rel.get("Id") for rel in rel_root}
    counter = 1
    while f"rIdSvg{counter}" in used_ids:
        counter += 1
    svg_rel_id = f"rIdSvg{counter}"
    svg_media_name = f"{svg_path.stem}.svg"
    svg_rel = etree.SubElement(rel_root, f"{{{relationship_ns}}}Relationship")
    svg_rel.set("Id", svg_rel_id)
    svg_rel.set("Type", image_rel_type)
    svg_rel.set("Target", f"media/{svg_media_name}")

    ext_list = blip.find(f"{{{drawing_ns}}}extLst")
    if ext_list is None:
        ext_list = etree.SubElement(blip, f"{{{drawing_ns}}}extLst")
    extension = etree.SubElement(ext_list, f"{{{drawing_ns}}}ext")
    extension.set("uri", "{96DAC541-7B7A-43D3-8B79-37D633B846F1}")
    svg_blip = etree.SubElement(extension, f"{{{svg_ns}}}svgBlip", nsmap={"asvg": svg_ns})
    svg_blip.set(f"{{{office_rel_ns}}}embed", svg_rel_id)

    content_types_path = "[Content_Types].xml"
    content_root = etree.fromstring(package[content_types_path])
    has_svg_type = any(
        item.get("Extension", "").lower() == "svg"
        for item in content_root.findall(f"{{{content_type_ns}}}Default")
    )
    if not has_svg_type:
        default = etree.SubElement(content_root, f"{{{content_type_ns}}}Default")
        default.set("Extension", "svg")
        default.set("ContentType", "image/svg+xml")

    package[rel_path] = etree.tostring(rel_root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    package[document_path] = etree.tostring(document_root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    package[content_types_path] = etree.tostring(content_root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    package[f"word/media/{svg_media_name}"] = svg_path.read_bytes()

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=docx_path.parent) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for name, payload in package.items():
                target.writestr(name, payload)
        temp_path.replace(docx_path)
    finally:
        temp_path.unlink(missing_ok=True)


def save_with_embedded_svg(doc: Document, output_path: Path, flow_name: str):
    save_with_embedded_svgs(doc, output_path, [flow_name])


def save_with_embedded_svgs(doc: Document, output_path: Path, flow_names: list[str]):
    doc.save(output_path)
    for flow_name in flow_names:
        _embed_svg_fallback(output_path, FLOWS / f"{flow_name}.svg", flow_name)


def _escape_svg(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_flow(name: str, title: str, rows: list[list[str]], terminal: str):
    """Build matching monochrome SVG and PNG flowcharts."""
    width = 1200
    margin = 55
    top = 95
    box_h = 82
    gap_y = 58
    max_cols = max(len(row) for row in rows)
    height = top + len(rows) * (box_h + gap_y) + 110
    svg = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#000"/></marker></defs>',
        '<rect width="100%" height="100%" fill="#fff"/>',
        f'<text x="{width/2}" y="42" text-anchor="middle" font-family="Arial" font-size="25" font-weight="700" fill="#000">{_escape_svg(title)}</text>',
    ]
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 25)
        box_font = ImageFont.truetype("arial.ttf", 18)
        box_font_b = ImageFont.truetype("arialbd.ttf", 18)
        small = ImageFont.truetype("arial.ttf", 15)
    except OSError:
        title_font = box_font = box_font_b = small = ImageFont.load_default()
    draw.text((width / 2, 28), title, fill="black", font=title_font, anchor="mm")

    centers: list[list[tuple[float, float]]] = []
    for r_idx, row in enumerate(rows):
        count = len(row)
        usable = width - 2 * margin
        gap_x = 38
        box_w = (usable - gap_x * (count - 1)) / count
        y = top + r_idx * (box_h + gap_y)
        current: list[tuple[float, float]] = []
        for c_idx, label in enumerate(row):
            x = margin + c_idx * (box_w + gap_x)
            svg.append(f'<rect x="{x:.1f}" y="{y}" width="{box_w:.1f}" height="{box_h}" rx="10" fill="#fff" stroke="#000" stroke-width="2"/>')
            lines = [part.strip() for part in label.split("\n")]
            text_y = y + box_h / 2 - (len(lines) - 1) * 11
            svg.append(f'<text x="{x + box_w/2:.1f}" y="{text_y:.1f}" text-anchor="middle" font-family="Arial" font-size="18" fill="#000">')
            for idx, line in enumerate(lines):
                dy = "0" if idx == 0 else "23"
                svg.append(f'<tspan x="{x + box_w/2:.1f}" dy="{dy}">{_escape_svg(line)}</tspan>')
            svg.append('</text>')
            draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=10, outline="black", width=2, fill="white")
            total_h = len(lines) * 22
            for idx, line in enumerate(lines):
                draw.text((x + box_w / 2, y + box_h / 2 - total_h / 2 + idx * 22 + 9), line, fill="black", font=box_font, anchor="mm")
            current.append((x + box_w / 2, y + box_h / 2))
        centers.append(current)

    for idx in range(len(centers) - 1):
        src_row, dst_row = centers[idx], centers[idx + 1]
        for src in src_row:
            targets = dst_row if len(src_row) == 1 else [min(dst_row, key=lambda d: abs(d[0] - src[0]))]
            for dst in targets:
                x1, y1 = src[0], src[1] + box_h / 2
                x2, y2 = dst[0], dst[1] - box_h / 2
                mid = (y1 + y2) / 2
                svg.append(f'<path d="M{x1:.1f},{y1:.1f} L{x1:.1f},{mid:.1f} L{x2:.1f},{mid:.1f} L{x2:.1f},{y2:.1f}" fill="none" stroke="#000" stroke-width="2" marker-end="url(#arrow)"/>')
                points = [(x1, y1), (x1, mid), (x2, mid), (x2, y2)]
                draw.line(points, fill="black", width=2)
                angle = math.atan2(y2 - mid, x2 - x2)
                draw.polygon([(x2, y2), (x2 - 7, y2 - 12), (x2 + 7, y2 - 12)], fill="black")

    end_y = top + len(rows) * (box_h + gap_y) - 20
    end_w = 560
    end_x = (width - end_w) / 2
    svg.append(f'<rect x="{end_x}" y="{end_y}" width="{end_w}" height="72" rx="36" fill="#000" stroke="#000"/>')
    svg.append(f'<text x="{width/2}" y="{end_y+44}" text-anchor="middle" font-family="Arial" font-size="19" font-weight="700" fill="#fff">{_escape_svg(terminal)}</text>')
    draw.rounded_rectangle((end_x, end_y, end_x + end_w, end_y + 72), radius=36, fill="black")
    draw.text((width / 2, end_y + 36), terminal, fill="white", font=box_font_b, anchor="mm")
    for src in centers[-1]:
        x1, y1 = src[0], src[1] + box_h / 2
        x2, y2 = width / 2, end_y
        mid = (y1 + y2) / 2
        svg.append(f'<path d="M{x1:.1f},{y1:.1f} L{x1:.1f},{mid:.1f} L{x2:.1f},{mid:.1f} L{x2:.1f},{y2:.1f}" fill="none" stroke="#000" stroke-width="2" marker-end="url(#arrow)"/>')
        draw.line([(x1, y1), (x1, mid), (x2, mid), (x2, y2)], fill="black", width=2)
        draw.polygon([(x2, y2), (x2 - 7, y2 - 12), (x2 + 7, y2 - 12)], fill="black")
    svg.append('</svg>')
    (FLOWS / f"{name}.svg").write_text("\n".join(svg), encoding="utf-8")
    img.save(ASSETS / f"{name}.png", "PNG")


def build_flows():
    build_flow(
        "core-platform-flow", "Argus Core Platform — Verified Data Flow",
        [["Authorized feeds\nand collectors", "Customer onboarding\nand recon", "Manual/API\ningest"],
         ["Normalize + deduplicate\nPostgreSQL", "Asset inventory\nand tenant context"],
         ["IOC extraction\n92 patterns / 17 families", "Direct correlation\nand threat pressure"],
         ["Enrichment + severity", "Attribution + campaigns", "Exposure scoring"]],
        "Findings · SLA · remediation · alerts · STIX / SIEM exports",
    )
    build_flow(
        "osint-investigation-flow", "OSINT Investigation — Verified Data Flow",
        [["Subject profile\nname / email / phone", "Username variants\nand site catalog"],
         ["Direct site checks", "SERP discovery", "CSV / Intelbase\nsearch"],
         ["HTTP + WAF-aware\nresponse analysis", "DOM / metadata\nverification"],
         ["Subject-match score", "Evidence-chain tier", "Cross-profile\ncorrelation"]],
        "Job result · confidence tier · evidence · pivots · export",
    )
    build_flow(
        "breach-correlation-flow", "Breach Search and Graph — Verified Data Flow",
        [["Email / phone / IP /\nname / username input", "Search filters +\npagination"],
         ["Type detection +\nnormalization", "Parameterized\nClickHouse queries"],
         ["Direct records", "Elasticsearch\ncorrelation", "Redis JSON cache"],
         ["Bounded entity BFS\nrecords ↔ PII", "Deduplication +\nstable SHA-256 IDs"]],
        "Search response · provenance · capped connection graph",
    )
    build_flow(
        "telegram-intelligence-flow", "Telegram Intelligence — Verified Data Flow",
        [["Authorized channels", "JSON / ZIP export", "Analyst text\ninput"],
         ["Normalize leetspeak\nand separators", "Extract links, channels\nand attachments"],
         ["Weighted category rules", "Saturating confidence", "Probabilistic\nrisk aggregation"],
         ["Persist message + IOC", "Customer asset match", "Queue downstream\npipeline"]],
        "Alerts · search / graph · findings · STIX 2.1 export",
    )
    build_flow(
        "pii-ioc-stix-flow", "PII, IOC and STIX — Verified Data Flow",
        [["Feed text / message /\nbreach record", "Customer assets"],
         ["Ordered pattern scan\n17 families", "PII normalization\nand entity extraction"],
         ["Confidence +\ndeduplication", "Exact / boundary / CIDR /\nCVE-product matching"],
         ["Detection + finding", "Deterministic STIX\nobject mapping"]],
        "Auditable evidence · customer context · interoperable export",
    )
    build_flow(
        "unified-platform-flow", "Argus Unified — End-to-End Walkthrough",
        [["Local bind / TLS ingress", "Built-in dashboard", "Built-in modules UI\n/modules"],
         ["FastAPI API + RBAC", "Mounted OSINT / Breach\nWSGI modules"],
         ["Collectors + recon", "Telegram intelligence", "Celery workers\nand scheduler"],
         ["PostgreSQL", "ClickHouse", "Elasticsearch", "Redis"],
         ["Detection → correlation", "Finding → SLA", "Exposure → report"]],
        "One operational platform · module truth · reproducible deployment",
    )


def _reference_style_flow(
    name: str,
    title: str,
    subtitle: str,
    layers: list[list[tuple[str, str, str]]],
    groups: list[tuple[int, int, str]],
    extra_edges: list[tuple[str, str, str, str]],
):
    """Render a dense, page-2-inspired grayscale architecture flow.

    Nodes are arranged in explicit runtime layers. The visual language mirrors
    the supplied reference: compact charcoal process nodes, decision diamonds,
    grouped execution regions, curved fan-out/fan-in connectors, and long
    safety/feedback paths. The canvas remains white for print readability.
    """

    width = 1400
    top = 140
    layer_gap = 148
    node_h = 80
    height = top + len(layers) * layer_gap + 96
    margin = 58
    gap_x = 26

    try:
        title_font = ImageFont.truetype("arialbd.ttf", 38)
        subtitle_font = ImageFont.truetype("arial.ttf", 18)
        node_font = ImageFont.truetype("arial.ttf", 17)
        node_bold = ImageFont.truetype("arialbd.ttf", 16)
        group_font = ImageFont.truetype("arialbd.ttf", 15)
        edge_font = ImageFont.truetype("arial.ttf", 13)
    except OSError:
        title_font = subtitle_font = node_font = node_bold = group_font = edge_font = ImageFont.load_default()

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((width / 2, 45), title, fill="black", font=title_font, anchor="mm")
    draw.text((width / 2, 85), subtitle, fill="#333333", font=subtitle_font, anchor="mm")
    draw.line((235, 112, width - 235, 112), fill="#9a9a9a", width=2)

    svg = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<defs><marker id="arrow" markerWidth="9" markerHeight="9" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#454545"/></marker></defs>',
        '<rect width="100%" height="100%" fill="#fff"/>',
        f'<text x="{width/2}" y="58" text-anchor="middle" font-family="Arial" font-size="38" font-weight="700" fill="#000">{_escape_svg(title)}</text>',
        f'<text x="{width/2}" y="91" text-anchor="middle" font-family="Arial" font-size="18" fill="#333">{_escape_svg(subtitle)}</text>',
        f'<line x1="235" y1="112" x2="{width-235}" y2="112" stroke="#999" stroke-width="2"/>',
    ]

    # Background execution regions.
    for start, end, label in groups:
        y = top + start * layer_gap - 24
        h = (end - start) * layer_gap + node_h + 56
        draw.rounded_rectangle((24, y, width - 24, y + h), radius=18, fill="#f7f7f7", outline="#8b8b8b", width=2)
        draw.rectangle((46, y - 1, 290, y + 26), fill="white")
        draw.text((58, y + 12), label.upper(), fill="#333333", font=group_font, anchor="lm")
        svg.extend([
            f'<rect x="24" y="{y}" width="{width-48}" height="{h}" rx="18" fill="#f7f7f7" stroke="#888" stroke-width="2"/>',
            f'<rect x="46" y="{y-1}" width="244" height="27" fill="#fff"/>',
            f'<text x="58" y="{y+18}" font-family="Arial" font-size="15" font-weight="700" fill="#333">{_escape_svg(label.upper())}</text>',
        ])

    node_map: dict[str, dict] = {}
    for layer_idx, layer in enumerate(layers):
        count = len(layer)
        usable = width - 2 * margin
        box_w = min(260, (usable - gap_x * (count - 1)) / count)
        total_w = box_w * count + gap_x * (count - 1)
        start_x = (width - total_w) / 2
        y = top + layer_idx * layer_gap
        for idx, (node_id, label, kind) in enumerate(layer):
            w = 168 if kind == "diamond" else box_w
            h = 120 if kind == "diamond" else node_h
            x = start_x + idx * (box_w + gap_x) + (box_w - w) / 2
            node_map[node_id] = {"x": x, "y": y, "w": w, "h": h, "label": label, "kind": kind}

    def anchor(node: dict, where: str):
        if where == "top":
            return node["x"] + node["w"] / 2, node["y"]
        if where == "bottom":
            return node["x"] + node["w"] / 2, node["y"] + node["h"]
        if where == "left":
            return node["x"], node["y"] + node["h"] / 2
        return node["x"] + node["w"], node["y"] + node["h"] / 2

    def bezier_points(p0, p1, p2, p3, steps=28):
        pts = []
        for i in range(steps + 1):
            t = i / steps
            q = 1 - t
            pts.append((
                q**3*p0[0] + 3*q*q*t*p1[0] + 3*q*t*t*p2[0] + t**3*p3[0],
                q**3*p0[1] + 3*q*q*t*p1[1] + 3*q*t*t*p2[1] + t**3*p3[1],
            ))
        return pts

    def draw_arrow(points):
        if len(points) < 2:
            return
        p1, p2 = points[-2], points[-1]
        angle = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
        size = 10
        left = (p2[0] - size * math.cos(angle - 0.55), p2[1] - size * math.sin(angle - 0.55))
        right = (p2[0] - size * math.cos(angle + 0.55), p2[1] - size * math.sin(angle + 0.55))
        draw.polygon([p2, left, right], fill="#454545")

    def draw_dashed_polyline(points, dash=11, gap=7):
        """Draw an exception/feedback connector without implying main flow."""
        for start, end in zip(points, points[1:]):
            dx, dy = end[0] - start[0], end[1] - start[1]
            length = math.hypot(dx, dy)
            if length == 0:
                continue
            ux, uy = dx / length, dy / length
            offset = 0.0
            while offset < length:
                segment_end = min(offset + dash, length)
                draw.line(
                    (
                        start[0] + ux * offset,
                        start[1] + uy * offset,
                        start[0] + ux * segment_end,
                        start[1] + uy * segment_end,
                    ),
                    fill="#555555",
                    width=2,
                )
                offset += dash + gap

    def wrap_node_label(label: str, font, max_width: float) -> list[str]:
        """Respect authored breaks and wrap any remaining overlong line."""
        wrapped: list[str] = []
        for authored_line in label.split("\n"):
            words = authored_line.split()
            if not words:
                wrapped.append("")
                continue
            current = words[0]
            for word in words[1:]:
                candidate = f"{current} {word}"
                if draw.textlength(candidate, font=font) <= max_width:
                    current = candidate
                else:
                    wrapped.append(current)
                    current = word
            wrapped.append(current)
        return wrapped

    def route_edge(src_id: str, dst_id: str, label: str = "", side: str = ""):
        src, dst = node_map[src_id], node_map[dst_id]
        if side in ("left", "right"):
            p0 = anchor(src, "left" if side == "left" else "right")
            p3 = anchor(dst, "left" if side == "left" else "right")
            lane = 34 if side == "left" else width - 34
            points = [p0, (lane, p0[1]), (lane, p3[1]), p3]
            draw_dashed_polyline(points)
            svg_path = f'M{p0[0]:.1f},{p0[1]:.1f} L{lane},{p0[1]:.1f} L{lane},{p3[1]:.1f} L{p3[0]:.1f},{p3[1]:.1f}'
            # Put the explanation beside the destination shape instead of on
            # the page edge, so the reader sees what the return path means.
            # Reserve a readable gutter for long return-path explanations.
            # Keeping the label at x=150 / x=width-150 avoids both the page
            # edge and same-layer process nodes.
            label_at = (95 if side == "left" else width - 95, p3[1] - 15)
            dash_attribute = ' stroke-dasharray="11 7"'
        else:
            p0 = anchor(src, "bottom")
            p3 = anchor(dst, "top")
            delta = max(34, abs(p3[1] - p0[1]) * 0.48)
            p1 = (p0[0], p0[1] + delta)
            p2 = (p3[0], p3[1] - delta)
            points = bezier_points(p0, p1, p2, p3)
            draw.line(points, fill="#555555", width=2)
            svg_path = f'M{p0[0]:.1f},{p0[1]:.1f} C{p1[0]:.1f},{p1[1]:.1f} {p2[0]:.1f},{p2[1]:.1f} {p3[0]:.1f},{p3[1]:.1f}'
            # Branch labels sit nearer their destination, which separates
            # sibling Yes/No paths that otherwise meet at the same midpoint.
            label_at = points[min(len(points) - 1, round(len(points) * 0.70))]
            dash_attribute = ""
        draw_arrow(points)
        svg.append(f'<path d="{svg_path}" fill="none" stroke="#555" stroke-width="2"{dash_attribute} marker-end="url(#arrow)"/>')
        if label:
            lx, ly = label_at
            box = draw.textbbox((lx, ly), label, font=edge_font, anchor="mm")
            label_w = max(70, box[2] - box[0] + 14)
            label_h = max(22, box[3] - box[1] + 8)
            draw.rounded_rectangle(
                (lx - label_w / 2, ly - label_h / 2, lx + label_w / 2, ly + label_h / 2),
                radius=4,
                fill="white",
                outline="#b5b5b5",
                width=1,
            )
            draw.text((lx, ly), label, fill="#222222", font=edge_font, anchor="mm")
            svg.append(f'<rect x="{lx-label_w/2:.1f}" y="{ly-label_h/2:.1f}" width="{label_w:.1f}" height="{label_h:.1f}" rx="4" fill="#fff" stroke="#b5b5b5" stroke-width="1"/>')
            svg.append(f'<text x="{lx:.1f}" y="{ly+5:.1f}" text-anchor="middle" font-family="Arial" font-size="13" fill="#222">{_escape_svg(label)}</text>')

    # Automatic adjacent-layer fan-out/fan-in edges.
    for idx in range(len(layers) - 1):
        current, nxt = layers[idx], layers[idx + 1]
        if len(current) == 1:
            pairs = [(current[0][0], n[0]) for n in nxt]
        elif len(nxt) == 1:
            pairs = [(n[0], nxt[0][0]) for n in current]
        else:
            pairs = []
            for pos, n in enumerate(current):
                target_idx = round(pos * (len(nxt) - 1) / max(1, len(current) - 1))
                pairs.append((n[0], nxt[target_idx][0]))
        for src, dst in pairs:
            route_edge(src, dst)

    for src, dst, label, side in extra_edges:
        route_edge(src, dst, label, side)

    # Nodes sit above connectors.
    for node in node_map.values():
        x, y, w, h, label, kind = node["x"], node["y"], node["w"], node["h"], node["label"], node["kind"]
        font = node_bold if kind in ("diamond", "terminal") else node_font
        font_size = 16 if kind in ("diamond", "terminal") else 17
        max_text_width = w * (0.63 if kind == "diamond" else 0.88)
        lines = wrap_node_label(label, font, max_text_width)
        if kind == "diamond":
            pts = [(x+w/2, y), (x+w, y+h/2), (x+w/2, y+h), (x, y+h/2)]
            draw.polygon(pts, fill="#4a4a4a", outline="black")
            svg.append(f'<polygon points="{x+w/2:.1f},{y} {x+w:.1f},{y+h/2:.1f} {x+w/2:.1f},{y+h:.1f} {x:.1f},{y+h/2:.1f}" fill="#4a4a4a" stroke="#000" stroke-width="2"/>')
            color = "white"
        elif kind == "control":
            draw.rounded_rectangle((x, y, x+w, y+h), radius=8, fill="#e8e8e8", outline="#222222", width=2)
            svg.append(f'<rect x="{x:.1f}" y="{y}" width="{w:.1f}" height="{h}" rx="8" fill="#e8e8e8" stroke="#222" stroke-width="2"/>')
            color = "black"
        elif kind == "store":
            draw.rounded_rectangle((x, y, x+w, y+h), radius=18, fill="white", outline="#111111", width=3)
            draw.line((x+10, y+15, x+w-10, y+15), fill="#777777", width=2)
            svg.extend([
                f'<rect x="{x:.1f}" y="{y}" width="{w:.1f}" height="{h}" rx="18" fill="#fff" stroke="#111" stroke-width="3"/>',
                f'<line x1="{x+10:.1f}" y1="{y+15}" x2="{x+w-10:.1f}" y2="{y+15}" stroke="#777" stroke-width="2"/>',
            ])
            color = "black"
        elif kind == "terminal":
            draw.rounded_rectangle((x, y, x+w, y+h), radius=h/2, fill="#111111", outline="black", width=2)
            svg.append(f'<rect x="{x:.1f}" y="{y}" width="{w:.1f}" height="{h}" rx="{h/2}" fill="#111" stroke="#000" stroke-width="2"/>')
            color = "white"
        else:
            draw.rounded_rectangle((x, y, x+w, y+h), radius=7, fill="#383838", outline="#111111", width=2)
            svg.append(f'<rect x="{x:.1f}" y="{y}" width="{w:.1f}" height="{h}" rx="7" fill="#383838" stroke="#111" stroke-width="2"/>')
            color = "white"
        line_step = 18 if kind == "diamond" else 20
        start_y = y + h/2 - (len(lines)-1)*line_step/2
        for line_idx, line in enumerate(lines):
            ty = start_y + line_idx*line_step
            draw.text((x+w/2, ty), line, fill=color, font=font, anchor="mm")
            svg.append(f'<text x="{x+w/2:.1f}" y="{ty+6:.1f}" text-anchor="middle" font-family="Arial" font-size="{font_size}" font-weight="{700 if kind in ("diamond", "terminal") else 400}" fill="#{"fff" if color == "white" else "000"}">{_escape_svg(line)}</text>')

    legend_y = height - 34
    legend = "Solid arrow: primary processing flow    |    Dashed labeled arrow: policy, exception, or feedback path    |    Diamond: decision"
    draw.text((width / 2, legend_y), legend, fill="#333333", font=edge_font, anchor="mm")
    svg.append(f'<text x="{width/2}" y="{legend_y+5}" text-anchor="middle" font-family="Arial" font-size="13" fill="#333">{_escape_svg(legend)}</text>')

    svg.append('</svg>')
    (FLOWS / f"{name}.svg").write_text("\n".join(svg), encoding="utf-8")
    img.save(ASSETS / f"{name}.png", "PNG")


# Override the earlier compact diagrams with detailed, reference-style maps.
def build_flows():
    _reference_style_flow(
        "core-platform-flow", "Core Threat Intelligence Flow", "Complete ingestion, correlation, scoring, and response runtime",
        [
            [("operator", "Analyst / Operator", "terminal"), ("safety", "Security Controls\nRBAC · Rate Limits", "control")],
            [("login", "Authenticate\nand Select Tenant", "normal"), ("scope", "Customer Scope\nand Owned Assets", "normal"), ("audit", "Audit / Metrics\nHealth / Kill Switch", "control")],
            [("gate", "Preflight\nGate", "diamond")],
            [("feeds", "Threat Feeds\nand Collectors", "normal"), ("recon", "Recon Engine\nDomains · Services", "normal"), ("api", "API / Manual\nEvent Ingest", "normal"), ("osint", "OSINT / Breach\nModule Inputs", "normal"), ("telegram", "Telegram\nAuthorized Inputs", "normal")],
            [("queue", "Celery Queue\nand Scheduler", "control"), ("ingest", "Unified Ingestion\nEnvelope", "normal"), ("raw", "Raw Evidence\nand Provenance", "store")],
            [("normalize", "Normalize · Hash\nDeduplicate", "normal"), ("patterns", "IOC Extraction\n92 Rules / 17 Families", "normal"), ("sourceconf", "Source Confidence\nand Time Decay", "normal")],
            [("classgate", "Evidence\nClass?", "diamond")],
            [("direct", "Direct Customer\nCorrelation", "normal"), ("pressure", "Environmental\nThreat Pressure", "normal"), ("global", "Global Unattributed\nIntel", "normal")],
            [("enrich", "Enrichment\nand Liveness", "normal"), ("severity", "Severity / SLA\nScoring", "normal"), ("attrib", "Actor / Campaign\nAttribution", "normal"), ("exposure", "Five-Dimension\nExposure Score", "normal")],
            [("finding", "Finding +\nMatch Proof", "normal"), ("action", "Remediation\nand Escalation", "normal"), ("history", "Exposure History\nand Trend", "store")],
            [("dashboard", "Dashboard / API\nAnalyst Review", "terminal"), ("exports", "STIX / SIEM\nReports / Alerts", "terminal")],
        ],
        [(0, 2, "Access and scope"), (3, 5, "Collection and normalization"), (6, 9, "Correlation and decision engine"), (10, 10, "Operational outputs")],
        [("safety", "gate", "Security policy check", "right"), ("audit", "dashboard", "Operational health", "right"), ("gate", "operator", "Blocked: return to operator", "left"), ("raw", "audit", "Evidence telemetry", "right"), ("global", "pressure", "Aggregate only", ""), ("finding", "patterns", "Detection feedback", "left")],
    )

    _reference_style_flow(
        "osint-investigation-flow", "OSINT Investigation Flow", "Profile-aware discovery, verification, correlation, and evidence control",
        [
            [("investigator", "Authorized\nInvestigator", "terminal"), ("policy", "Scope / Privacy\nRate Policy", "control")],
            [("profile", "Subject Profile\nNames · Email · Phone", "normal"), ("job", "Job Controller\nStart · Pause · Resume", "control")],
            [("scopegate", "Authorized\nScope?", "diamond")],
            [("perms", "Username / Alias\nPermutations", "normal"), ("sites", "Site Catalog +\nDirect Checks", "normal"), ("serp", "SERP Discovery\nDork Queries", "normal"), ("csv", "CSV / Intelbase\nLocal Search", "normal")],
            [("http", "Bounded HTTP\nRetries · Delays", "normal"), ("status", "Status / Redirect\nSignals", "normal"), ("waf", "WAF / Challenge\nDetection", "diamond")],
            [("markers", "Positive / Negative\nMarkers", "normal"), ("dom", "DOM Skeleton\nJaccard", "normal"), ("meta", "Profile Metadata\nBio · Links · Location", "normal"), ("alt", "Alternate JSON /\nPublic Endpoints", "normal")],
            [("match", "Subject-Match\nScore", "normal"), ("strength", "Input Profile\nStrength Prior", "normal"), ("evidence", "Evidence Chain\nSignal Weights", "normal")],
            [("tiergate", "Confidence\nTier", "diamond")],
            [("confirmed", "Confirmed / High\nCandidates", "normal"), ("ambiguous", "Ambiguous /\nUnverified", "normal"), ("notfound", "Not Found /\nNegative Evidence", "normal")],
            [("correlate", "Cross-Platform\nCorrelation", "normal"), ("pivots", "New Anchors\nEmails · Phones · URLs", "normal"), ("calibrate", "Known +/- Account\nCalibration", "control")],
            [("result", "Result + Evidence\nProgress / Export", "terminal"), ("review", "Analyst Review\nand Attribution Decision", "terminal")],
        ],
        [(0, 2, "Authority and subject setup"), (3, 5, "Discovery and response verification"), (6, 9, "Scoring and disambiguation"), (10, 10, "Evidence outputs")],
        [("policy", "scopegate", "Scope + privacy rules", "right"), ("scopegate", "investigator", "Blocked: revise scope", "left"), ("waf", "job", "Challenge: retry or pause", "right"), ("ambiguous", "profile", "Refine subject profile", "left"), ("calibrate", "sites", "Site calibration", "right"), ("review", "pivots", "Analyst-approved pivot", "right")],
    )

    _reference_style_flow(
        "breach-correlation-flow", "Breach Search and Graph Flow", "Exact lookup, cached correlation, bounded PII graph expansion, and provenance",
        [
            [("analyst", "Authorized Analyst", "terminal"), ("privacy", "Dataset Policy\nRBAC · Audit", "control")],
            [("seed", "Search Seed\nEmail · Phone · Name", "normal"), ("filters", "Field Filters\nLimit · Offset", "normal")],
            [("validate", "Valid Input\nand Field?", "diamond")],
            [("type", "PII Type\nDetection", "normal"), ("norm", "Normalize Email /\nPhone / Name", "normal"), ("catalog", "Controlled Field\nCatalog", "control")],
            [("cachegate", "Redis Cache\nHit?", "diamond")],
            [("click", "Parameterized\nClickHouse Search", "normal"), ("elastic", "Elasticsearch\nCorrelation Pivots", "normal"), ("cached", "Cached JSON\nResponse", "store")],
            [("records", "Deduplicated Records\nSource Provenance", "normal"), ("count", "Count + Pagination\nMetadata", "normal")],
            [("entities", "Extract PII Entities\nEmail · Phone · Name · Address", "normal"), ("ids", "Stable SHA-256\nNode / Edge IDs", "normal")],
            [("bfs", "Bounded BFS Queue\nBatch Size ≤ 10", "normal"), ("common", "Common Value\n> 200 Records?", "diamond"), ("caps", "Graph Caps\n500 Records / 300 Entities", "diamond")],
            [("expand", "Expand First-Degree\nRecords Only", "normal"), ("sample", "Retain 10 Samples\nMark Warning", "control"), ("stop", "Stop + Cap Reason", "control")],
            [("response", "Search Response\nand Connection Graph", "terminal"), ("review", "Analyst Review\nNo Identity Assumption", "terminal")],
        ],
        [(0, 2, "Access and query contract"), (3, 6, "Search execution and cache"), (7, 9, "Bounded relationship graph"), (10, 10, "Evidence outputs")],
        [("privacy", "validate", "Dataset access rules", "right"), ("validate", "analyst", "Rejected: correct the query", "left"), ("cached", "records", "Cache hit: reuse response", ""), ("common", "sample", "Yes: keep 10 samples", ""), ("common", "expand", "No: continue expansion", ""), ("caps", "stop", "Limit reached: stop", ""), ("response", "cachegate", "Store completed response", "right")],
    )

    _reference_style_flow(
        "telegram-intelligence-flow", "Telegram Threat Intelligence Flow", "Authorized collection/import, risk classification, IOC routing, and export",
        [
            [("operator", "Analyst / Operator", "terminal"), ("guard", "Authorization /\nChannel Allowlist", "control")],
            [("live", "Telethon Session\nOne-Time Authorization", "normal"), ("import", "JSON / ZIP\nAuthorized Export", "normal"), ("analyze", "Stateless Text\nAnalyze", "normal")],
            [("session", "Session Valid\nand Authorized?", "diamond"), ("bounds", "Import Bounds\n50 MiB / 10k Records", "diamond")],
            [("channels", "Explicit Channels\nOnly", "normal"), ("messages", "Message + Sender\nMedia Metadata", "normal"), ("audit", "Import Run\nAudit Record", "store")],
            [("normalize", "Leetspeak / Separator\nNormalization", "normal"), ("extract", "Links · Mentions\nAttachments · IOCs", "normal")],
            [("ransom", "Ransomware\nRules", "normal"), ("breach", "Data Breach /\nInitial Access", "normal"), ("fraud", "Carding /\nPhishing", "normal"), ("exploit", "Exploit / Broker\nRules", "normal")],
            [("confidence", "Saturating Category\nConfidence", "normal"), ("risk", "Probabilistic Risk\nComplement Boosts", "normal")],
            [("flag", "Flagged?\nRisk ≥ 55", "diamond")],
            [("persist", "PostgreSQL Message\nChannel / IOC", "store"), ("alerts", "Alert Feed +\nRisk Reasons", "normal"), ("search", "Exact / Fuzzy\nSearch + Graph", "normal")],
            [("match", "Customer Asset\nIOC Correlation", "normal"), ("finding", "Finding / SLA /\nRemediation Pipeline", "normal"), ("stix", "STIX 2.1\nBundle Mapping", "normal")],
            [("dashboard", "Telegram Workspace\nand Analyst Review", "terminal"), ("exports", "Alerts · Findings\nSTIX Export", "terminal")],
        ],
        [(0, 2, "Authorization and inputs"), (3, 6, "Collection and deterministic classification"), (7, 9, "Decision and platform integration"), (10, 10, "Analyst outputs")],
        [("guard", "session", "Allowlist + auth rules", "right"), ("session", "operator", "Unauthorized: re-authorize", "left"), ("bounds", "audit", "Bounds failure: rollback", "right"), ("flag", "alerts", "Yes: create alert", ""), ("flag", "persist", "Always preserve evidence", ""), ("finding", "confidence", "Pipeline feedback", "left"), ("dashboard", "search", "Analyst search query", "right")],
    )

    _reference_style_flow(
        "pii-ioc-stix-flow", "PII, IOC and STIX Flow", "Ordered extraction, cautious attribution, customer context, and interoperable output",
        [
            [("input", "Feed / Message /\nBreach Evidence", "terminal"), ("policy", "Sensitive Data\nAccess Policy", "control")],
            [("scan", "Ordered Pattern Scan\nSpecific → Generic", "normal"), ("context", "Customer Assets\nand Tenant Context", "normal")],
            [("valid", "Candidate Passes\nNoise Guards?", "diamond")],
            [("creds", "Credentials / Keys\nTokens / Sessions", "normal"), ("network", "Network / Domain\nEmail / Hash", "normal"), ("pii", "Financial / Identity\nPII Shapes", "normal"), ("actor", "Actor / Exfil /\nCVE / Crypto", "normal")],
            [("normalize", "Normalize Value\nand Context", "normal"), ("dedupe", "Deduplicate by\nCategory · Type · Value", "normal"), ("confidence", "Base Confidence +\nDomain Context Boost", "normal")],
            [("attribgate", "Attributable to\nCustomer?", "diamond")],
            [("ip", "Exact IP / CIDR\nMembership", "normal"), ("domain", "Exact / Subdomain\nBoundary", "normal"), ("credential", "Email Domain /\nExecutive Match", "normal"), ("cve", "CVE → Product →\nVersion Check", "normal"), ("brand", "Brand / Keyword\nEvidence", "normal")],
            [("direct", "Customer Detection\nand Match Proof", "normal"), ("probable", "Probable Exposure\nUnknown Version", "normal"), ("global", "Global Threat Signal\nNo False Attribution", "normal")],
            [("stixgate", "STIX Mapping\nType?", "diamond")],
            [("standard", "Standard Observable\nSTIX Pattern", "normal"), ("vuln", "CVE Vulnerability\nObject", "normal"), ("custom", "Open pattern_type\nargus", "normal")],
            [("bundle", "Deterministic UUIDv5\nDeduplicated Bundle", "terminal"), ("review", "Detection / Finding\nAnalyst Review", "terminal")],
        ],
        [(0, 2, "Evidence intake and validation"), (3, 4, "IOC and PII extraction"), (5, 7, "Cautious customer correlation"), (8, 10, "STIX interoperability and review")],
        [("policy", "valid", "Sensitive-data rules", "right"), ("valid", "input", "Rejected as noise", "left"), ("attribgate", "global", "No: global signal only", ""), ("attribgate", "ip", "Yes: test customer assets", ""), ("probable", "review", "Review at lower confidence", "right"), ("review", "context", "Customer-asset feedback", "right"), ("bundle", "scan", "STIX validation feedback", "left")],
    )

    _reference_style_flow(
        "unified-platform-flow", "Argus Unified Complete Flow", "End-to-end product architecture and runtime across all integrated modules",
        [
            [("user", "Analyst / Administrator", "terminal"), ("ops", "Operations / Policy\nSecrets · TLS · Backups", "control")],
            [("edge", "Local Bind or\nExternal TLS Ingress", "normal"), ("auth", "JWT / RBAC\nTenant Context", "normal")],
            [("ui", "Built-in Analyst\nDashboard", "normal"), ("modulesui", "Built-in Modules SPA\n/ modules", "normal"), ("api", "FastAPI OpenAPI\n170 Operations", "normal")],
            [("osint", "OSINT Investigation\nMounted /api/osint", "normal"), ("breach", "Breach Search / Graph\nMounted /api/breach", "normal"), ("telegram", "Telegram Intelligence\n/api/telegram", "normal"), ("core", "Core CTI / Customer\nFindings / Exposure", "normal")],
            [("inputgate", "Authorized Input\nand Scope?", "diamond")],
            [("collect", "Collectors / Recon\nManual / API Ingest", "normal"), ("jobs", "Celery Workers\nand Scheduler", "control"), ("imports", "OSINT / Breach /\nTelegram Inputs", "normal")],
            [("postgres", "PostgreSQL\nSystem of Record", "store"), ("redis", "Redis\nQueue + Cache", "store"), ("clickhouse", "ClickHouse\nBreach Corpus", "store"), ("elastic", "Elasticsearch\nCorrelation", "store")],
            [("normalize", "Normalize / Deduplicate\nPreserve Provenance", "normal"), ("detect", "IOC / PII / Threat\nDetection", "normal"), ("correlate", "Customer Asset\nCorrelation", "normal")],
            [("decision", "Direct Finding or\nThreat Pressure?", "diamond")],
            [("enrich", "Enrichment / Severity\nAttribution / Campaign", "normal"), ("finding", "Finding / SLA /\nRemediation", "normal"), ("exposure", "Exposure Score /\nHistory / Narrative", "normal")],
            [("dashboard", "Unified Analyst Workspace\nSearch · Review · Operate", "terminal"), ("output", "STIX · SIEM · Alerts\nReports · Metrics", "terminal")],
        ],
        [(0, 2, "Single Argus image, identity, and user experience"), (3, 5, "Integrated module execution"), (6, 8, "Shared data and correlation plane"), (9, 10, "Operational intelligence outputs")],
        [("ops", "inputgate", "Operational guardrails", "right"), ("inputgate", "user", "Blocked: correct scope", "left"), ("jobs", "redis", "Queue and retry state", ""), ("postgres", "api", "Shared application models", "right"), ("dashboard", "ui", "Analyst actions", "left"), ("output", "ops", "Audit and monitoring", "right"), ("finding", "detect", "Detection feedback", "left")],
    )


def core_document():
    doc = new_document("Core Platform")
    masthead(doc, "Core Threat Intelligence Platform", "Canonical ArgusWatch shell, shared services, scoring, correlation, operations, and exports")
    doc.add_heading("Executive implementation summary", level=1)
    add_para(doc, "Argus Unified keeps argus_main as the canonical product shell. One multi-stage Argus image contains the FastAPI service, built-in dashboard, compiled modules SPA, workers, breach roles, intelligence proxy, and scoped recon engine. PostgreSQL, Redis, ClickHouse, and Elasticsearch remain independent infrastructure services. OSINT, breach, Telegram, and PII/IOC capabilities are modules inside the shared application rather than parallel products.")
    add_para(doc, "The OpenAPI schema exposes 156 native FastAPI paths and 170 HTTP operations. The legacy OSINT and breach applications are mounted under /api/osint and /api/breach, avoiding the previous broad /api path collision.")
    add_lineage(doc, [SOURCE_DOCS[0], "Architecture and operational flow were retained only where verified against the consolidated repository."])

    doc.add_heading("Implemented architecture", level=1)
    add_flow(doc, "core-platform-flow", "Figure 1. Core ingestion, correlation, scoring, and response path.")
    add_table(doc, ["Layer", "Implemented components", "Responsibility"], [
        ["Entry/UI", "Local bind or external TLS ingress; built-in dashboard and modules SPA", "Access policy, static delivery, and analyst workflows"],
        ["API", "FastAPI + mounted a2wsgi Flask modules", "RBAC, validation, orchestration, module endpoints"],
        ["Execution", "Celery workers and beat", "Collection, enrichment, correlation, exports, scheduled jobs"],
        ["State", "PostgreSQL, Redis, ClickHouse, Elasticsearch", "System of record, cache/queue, breach data, search correlation"],
        ["Collection", "Native collectors, recon engine, Telegram collector", "Authorized source acquisition and normalization"],
        ["Outputs", "Findings, SLA, reports, STIX, SIEM/syslog", "Actionable analyst and integration artifacts"],
    ], [1300, 3350, 4710])

    doc.add_heading("Primary workflows", level=2)
    add_numbered(doc, [
        "Create a customer, set sector metadata, and register owned domains, IPs, CIDRs, brands, executives, and technology-stack assets.",
        "Run authorized collectors, recon, Telegram import/collection, or explicit API ingestion. Inputs become normalized detections or module records.",
        "Extract IOCs, suppress duplicates, apply source confidence and freshness, and preserve raw evidence/provenance.",
        "Route direct evidence to customers using exact IP, CIDR, domain-boundary, credential-domain, brand, and CVE-to-product/version strategies.",
        "Convert unmatched activity into environmental threat-pressure signals rather than forcing false customer attribution.",
        "Enrich and score findings, attach attribution/campaign context, generate remediation and SLA state, and expose results through the UI/API.",
    ])

    add_core_math(doc)

    doc.add_heading("Matching correctness controls", level=2)
    add_bullets(doc, [
        "Domain comparison uses exact host, subdomain boundary, or delimited occurrence; raw substring matching is not used.",
        "CIDR matching uses Python ipaddress with strict=False and skips invalid networks.",
        "CVE matches pass through product normalization and affected-version checks. Unknown versions create lower-confidence probable exposures instead of confirmed findings.",
        "Cross-source confirmation can raise severity by one level; provenance and match proof are stored on findings.",
        "AI automation is opt-in and disabled by default. Deterministic collection, matching, scoring, and reporting do not require an AI provider.",
    ])

    doc.add_heading("Security and operations", level=1)
    add_bullets(doc, [
        "Authentication fails closed unless AUTH_DISABLED is explicitly set for controlled development tests.",
        "JWT handling uses PyJWT with cryptographic extras; API writes are role-gated and rate limiting is enabled.",
        "Only the FastAPI entry port is published, and it binds to 127.0.0.1 by default. Databases, queues, search services, collectors, and recon remain internal to Compose.",
        "Required secrets have no production fallback in Compose. Production TLS terminates at a reviewed reverse proxy, load balancer, VPN, or cluster ingress.",
        "The API, workers, breach roles, intelligence proxy, recon engine, dashboard, and modules SPA are built once and reuse the same immutable Argus image.",
        "Migrations are idempotent and consolidate the duplicate dark-web table into dark_web_mentions.",
        "Start, stop, backup/migration, and one-time Telegram authorization commands are documented in README.md.",
    ])

    add_comparison(doc, *CORE_COMPARISON)
    add_para(doc, "Selection context: OpenCTI and MISP are stronger fits when the primary requirement is a mature CTI knowledge-sharing ecosystem. ThreatConnect adds commercial intelligence and case APIs. Recorded Future supplies a large managed intelligence corpus and attack-surface intelligence. Argus is differentiated here by combining customer-owned assets, recon, breach search, Telegram analysis, deterministic exposure scoring, findings, SLA, and remediation in one self-hosted implementation; it does not include a commercial intelligence corpus.")
    add_sources(doc, ["nist_risk", "temporal_ir", "opencti", "opencti_connectors", "misp", "threatconnect_indicators", "threatconnect_associations", "threatconnect_cases", "threatconnect_endpoints", "threatconnect_roles", "threatconnect_stix", "recorded_future"])
    add_verification(doc, ["OpenAPI inspection: 156 native paths and 170 operations; mounted legacy modules are intentionally outside the generated FastAPI schema.", "Docker Compose configuration resolves with required test secrets. Full container runtime validation was unavailable because the local Docker Linux engine was not running."])
    add_boundaries(doc, [
        "Live collector behavior depends on source credentials, quotas, availability, and program authorization.",
        "Enterprise collector status endpoints include configured adapters and stubs; a stub is not documented as a working data source.",
        "A real PostgreSQL/ClickHouse/Elasticsearch/Redis container pass remains required on a host with a running Docker Linux engine before production promotion.",
    ])
    save_with_embedded_svg(doc, MODULES / "01_Argus_Core_Platform.docx", "core-platform-flow")


def osint_document():
    doc = new_document("OSINT Investigation")
    masthead(doc, "OSINT Investigation Module", "Profile-aware discovery, verification, evidence chains, correlation, and resumable job control")
    doc.add_heading("Implemented scope", level=1)
    add_para(doc, "The module accepts a username or a richer subject profile and runs direct site checks, candidate generation, search-engine discovery, CSV/Intelbase lookup, profile parsing, evidence scoring, cross-platform correlation, and export. Synchronous and background workflows are mounted at /api/osint; background investigations support status, result, pause, resume, cancel, and bounded progress state.")
    add_lineage(doc, [SOURCE_DOCS[2], "Verified concepts reused: candidate generation, multi-signal verification, calibration, DOM comparison, evidence tiers, and job controls."])
    add_flow(doc, "osint-investigation-flow", "Figure 1. Profile-aware OSINT investigation and verification path.")

    doc.add_heading("Endpoints and inputs", level=2)
    add_table(doc, ["Endpoint family", "Implemented behavior", "Representative input"], [
        ["/api/osint/search*", "Direct username checks; start/status/result/cancel; recursive search", "username, mode, site filters"],
        ["/api/osint/v2/investigate*", "Profile-aware synchronous/background investigation with pause/resume", "name, aliases, email, phone, location, employment, education"],
        ["/api/osint/v2/quick", "Reduced-latency profile investigation", "subject profile"],
        ["/api/osint/v2/calibrate", "Known-positive/known-negative detection calibration", "optional site list"],
        ["/api/osint/v2/export", "Structured result export", "completed investigation payload"],
        ["/api/osint/csv-only", "Local authorized CSV search path", "identifier and selected fields"],
    ], [2250, 3950, 3160])

    add_osint_math(doc)

    doc.add_heading("Reliability controls", level=1)
    add_bullets(doc, [
        "Site-specific positive/negative markers, redirect behavior, WAF/challenge detection, response status, JSON markers, and page metadata are evaluated together.",
        "HTTP concurrency, retries, delays, per-site configuration, cancellation, and progress callbacks bound active investigations.",
        "Calibration checks known-existing and known-nonexistent accounts to expose broken site rules.",
        "Correlation adds bounded boosts for usernames, names, dates of birth, cross-links, and metadata richness while preserving source evidence.",
    ])

    add_comparison(doc, *OSINT_COMPARISON)
    add_para(doc, "Selection context: SpiderFoot is broader for automated infrastructure and OSINT module coverage; Maltego is stronger for interactive graph analysis, transforms, and team collaboration; Sherlock is a focused username enumerator; Maigret provides deep username dossiers and many export formats. Argus focuses on profile-aware identity disambiguation, explicit evidence tiers, DOM/marker verification, calibration, resumable jobs, and direct handoff into breach and threat workflows.")
    add_sources(doc, ["fellegi_sunter", "jaccard", "spiderfoot", "maltego", "maltego_transforms", "maltego_collaboration", "sherlock", "maigret"])
    add_verification(doc, ["OSINT scoring, type detection, verification helpers, and API validation are covered by the consolidated backend tests."])
    add_boundaries(doc, [
        "A positive site response is evidence, not identity proof. Analyst review is required for high-impact decisions.",
        "Coverage changes as sites modify markup, authentication, rate limits, or bot defenses; calibration exposes but cannot prevent external drift.",
        "The module does not bypass authentication controls or access private profiles. Use must be authorized and lawful.",
    ])
    save_with_embedded_svg(doc, MODULES / "02_OSINT_Investigation_Module.docx", "osint-investigation-flow")


def breach_document():
    doc = new_document("Breach Search and Graph")
    masthead(doc, "Breach Search and Correlation Module", "Direct ClickHouse search, Elasticsearch pivots, Redis caching, and bounded PII relationship graphs")
    doc.add_heading("Implemented scope", level=1)
    add_para(doc, "The breach module is mounted at /api/breach for the product UI and also contains its canonical FastAPI search service. It performs parameterized direct searches over ClickHouse, optional Elasticsearch correlation, Redis JSON caching, combined responses, and a bounded breadth-first connection graph. It does not claim a proprietary breach corpus; results depend on authorized datasets loaded into the configured stores.")
    add_lineage(doc, [SOURCE_DOCS[3], "Verified concepts reused: ClickHouse exact search, Elasticsearch correlation, Redis cache, field catalog, provenance, and bounded graph expansion."])
    add_flow(doc, "breach-correlation-flow", "Figure 1. Direct breach lookup, correlation, caching, and graph expansion.")

    doc.add_heading("Search interfaces", level=2)
    add_table(doc, ["Interface", "Behavior", "Bounds"], [
        ["POST /api/breach/search", "UI gateway to the canonical search service", "limit 1–500; non-negative offset"],
        ["POST /api/breach/graph/connections", "Connection graph from email, phone, or name seed", "timeout and hard graph caps"],
        ["GET /search", "Field/value direct ClickHouse query", "default 50; maximum 500"],
        ["POST /search/direct", "Multi-filter exact/direct search", "server-enforced maximum"],
        ["POST /search/correlation", "Elasticsearch pivot search", "max results per query"],
        ["POST /search/combined", "Direct and correlated results in one response", "independent direct offsets"],
        ["POST /search/connections", "Canonical bounded graph", "500 records + 300 entities"],
    ], [2750, 4020, 2590])

    doc.add_heading("Normalization and deterministic identifiers", level=1)
    add_bullets(doc, [
        "Input type detection recognizes email, phone, IPv4, Aadhaar-shaped values, names, usernames, and unknown input.",
        "Emails are trimmed and lowercased. Phone numbers remove non-digits and normalize common country/trunk prefixes before retaining a ten-digit form for search.",
        "Direct filters are compiled into parameterized ClickHouse predicates; field names come from a controlled field catalog rather than raw client SQL.",
        "Optional privacy-preserving email indexing uses SHA-256(salt + ':' + normalized_email).",
    ])
    add_code(doc, "row_id = SHA256(source_file | lowercase(email) | normalized_phone)\nentity_id = SHA256(UPPER(entity_type) : lowercase(trim(value)))\nedge_id = SHA256(row_id | entity_id)")

    doc.add_heading("Bounded connection graph", level=1)
    add_numbered(doc, [
        "Resolve up to 100 seed records using the normalized email, phone, or name.",
        "Extract EMAIL, PHONE, NAME, and ADDRESS entities from each record and enqueue unvisited entities.",
        "Query entities in concurrent batches of at most ten, requesting 201 records to detect common values.",
        "When an entity returns more than 200 records, mark it common and retain only ten samples instead of expanding it.",
        "Expand only first-degree records into second-degree entities; deduplicate all nodes and edges by stable hashes.",
        "Stop at the combined 800-node ceiling (500 records plus 300 entities) and return capped/cap_reason metadata.",
    ])
    add_para(doc, "The algorithm is breadth-first over a bipartite graph of breach records and extracted entities. The hard caps and common-value short circuit prevent names or shared addresses from producing unbounded traversal.")

    doc.add_heading("Data and failure behavior", level=1)
    add_bullets(doc, [
        "ClickHouse is authoritative for direct bulk breach records; Elasticsearch supplies correlation pivots; Redis cache failures degrade to uncached queries.",
        "Every returned record preserves source_file provenance where present.",
        "Count queries and data queries are separate; pagination returns limit, offset, and total metadata.",
        "Search errors are converted to bounded API errors without exposing database credentials or raw stack traces.",
    ])

    add_breach_math(doc)

    add_comparison(doc, *BREACH_COMPARISON)
    add_para(doc, "Selection context: HIBP is optimized for breach membership, verified-domain exposure, stealer-log and notification use cases without exposing raw breach rows. Intelligence X, LeakCheck, and DeHashed provide vendor-hosted search corpora with different field access and commercial tiers. Argus instead searches operator-supplied authorized stores and adds a bounded record-to-PII graph plus customer finding correlation; it ships with no breach dataset.")
    add_sources(doc, ["moore_bfs", "meta_blocking", "fips180", "hibp", "intelx", "leakcheck", "dehashed"])
    add_verification(doc, ["Input classification, normalization, query construction, graph caps, and mocked graph traversal are exercised without requiring production breach data."])
    add_boundaries(doc, [
        "No breach records are bundled. The operator must lawfully obtain, secure, retain, and delete datasets under applicable policy and law.",
        "Data-store end-to-end validation requires running ClickHouse, Elasticsearch, and Redis services and representative authorized fixtures.",
        "Name/address correlation is heuristic and can join different people; evidence must be reviewed before action.",
    ])
    save_with_embedded_svg(doc, MODULES / "03_Breach_Search_and_Graph_Module.docx", "breach-correlation-flow")


def telegram_document():
    doc = new_document("Telegram Threat Intelligence")
    masthead(doc, "Telegram Threat Intelligence Module", "Authorized collection/import, obfuscation-aware classification, fuzzy search, IOC routing, and STIX export")
    doc.add_heading("Implemented scope", level=1)
    add_para(doc, "This is the unique functionality consolidated from cyber_threat_monitor and pii_link into the shared Argus platform. It supports one-time Telethon session authorization, explicit-channel collection, JSON/ZIP imports, deterministic text analysis, persistent message/channel/import audit records, typo-tolerant search, alerts, graph summaries, customer asset correlation, downstream finding creation, and STIX 2.1 export.")
    add_lineage(doc, [SOURCE_DOCS[1], SOURCE_DOCS[4], "The old SQLite/Elasticsearch/Neo4j service stacks were not duplicated. Their verified Telegram classification and entity-correlation concepts now use Argus PostgreSQL, Redis/Celery, IOC matching, findings, and exports."])
    add_flow(doc, "telegram-intelligence-flow", "Figure 1. Telegram ingestion, classification, correlation, and response path.")

    doc.add_heading("API and storage", level=1)
    add_table(doc, ["Endpoint", "Implemented behavior"], [
        ["GET /api/telegram/health", "Stateless module health"],
        ["POST /api/telegram/analyze", "Offline deterministic analysis without persistence"],
        ["POST/GET /api/telegram/messages", "Persist and list messages; exact or fuzzy query"],
        ["GET /api/telegram/alerts", "Recent flagged messages ordered by risk and time"],
        ["GET /api/telegram/stats", "Message, flagged, channel, category, and import counts"],
        ["GET /api/telegram/channels", "Channel activity summary"],
        ["GET /api/telegram/graph", "Channel/reference relationship data"],
        ["GET /api/telegram/stix/export", "STIX bundle for selected/recent messages"],
        ["POST /api/telegram/import", "Bounded JSON/ZIP import with audit and rollback"],
    ], [3150, 6210])
    add_para(doc, "Persistence uses telegram_messages, telegram_channels, and telegram_import_runs in the shared PostgreSQL schema. Imported or collected messages are processed by the same IOC matcher and finding pipeline used by other intelligence sources.")

    add_telegram_math(doc)

    doc.add_heading("Import safety and collection model", level=1)
    add_bullets(doc, [
        "Upload limit: 50 MiB; expanded ZIP limit: 250 MiB; maximum 2,000 archive files; maximum 10,000 imported records.",
        "A failed import rolls back partial message writes while retaining the failed import-run audit record.",
        "Live collection reads only explicitly configured channels and bounded textual attachment types.",
        "The Celery worker is non-interactive. If its persisted session is not authorized, it returns session_not_authorized; authorization is performed once with scripts/authorize_telegram.py.",
        "Telegram API ID/hash and user authorization are required; no credentials or session are bundled.",
    ])

    add_comparison(doc, *TELEGRAM_COMPARISON)
    add_para(doc, "Selection context: Telethon is the flexible MTProto client library used by Argus, not an intelligence product by itself. Maltego Monitor provides paid continuous public-source monitoring and filtering. Social Links provides commercial Telegram transforms and relationship investigation, with advanced Telegram packs restricted to eligible government customers. OpenCTI supplies a CTI graph, cases, alerts, and STIX connectors but no native Telegram collector was evidenced. Argus supplies the Telegram-specific classification, risk mathematics, bounded import, customer matching, and STIX pipeline.")
    add_sources(doc, ["noisy_or", "ratcliff", "telethon", "telethon_events", "maltego_monitor", "maltego_telegram", "sociallinks", "sociallinks_telegram", "opencti", "opencti_connectors"])
    add_verification(doc, ["Telegram detector, fuzzy matcher, import bounds, STIX mapping, and API analysis inputs are covered by automated tests.", "Browser input verification confirmed ransomware text, CVE extraction, Telegram URL extraction, 78 risk, 91% ransomware confidence, and flagged output in the final UI."])
    add_boundaries(doc, [
        "Live MTProto collection was not executed because no user Telegram credentials/session were supplied.",
        "Classification is deterministic lexical evidence, not authorship proof or a substitute for analyst validation.",
        "Only channels the operator is authorized to access may be configured.",
    ])
    save_with_embedded_svg(doc, MODULES / "04_Telegram_Threat_Intelligence_Module.docx", "telegram-intelligence-flow")


def pii_document():
    doc = new_document("PII, IOC and STIX")
    masthead(doc, "PII, IOC Correlation and STIX Module", "Ordered extraction, confidence, entity normalization, customer correlation, and interoperable CTI output")
    doc.add_heading("Implemented scope", level=1)
    add_para(doc, "The original pii_link concepts are integrated as shared platform capabilities rather than a duplicate standalone service. Argus scans unstructured intelligence, breach records, and Telegram messages for structured indicators; normalizes PII entities; correlates attributable evidence to customer assets; preserves confidence/provenance; and emits STIX 2.1-compatible bundles.")
    add_lineage(doc, [SOURCE_DOCS[4], SOURCE_DOCS[0], "The corrupted embedded image in the original pii.docx did not prevent reuse of its document XML text. Neo4j and a second Elasticsearch/Redis stack were deliberately excluded because equivalent shared services already existed."])
    add_flow(doc, "pii-ioc-stix-flow", "Figure 1. Shared PII/IOC extraction, customer correlation, and STIX mapping.")

    doc.add_heading("IOC extraction", level=1)
    add_para(doc, "The ordered scanner implements 92 regex patterns across 17 families. Specific families run before generic hashes, URLs, domains, emails, and IPs so a narrow token is not stolen by a broad pattern. Results are deduplicated by category, IOC type, and exact value, then sorted by confidence.")
    add_table(doc, ["Pattern families", "Representative outputs", "Reliability guard"], [
        ["Credentials / API keys / OAuth / sessions", "email-password pairs, AWS/GitHub/Slack/Google tokens, JWT, NTLM", "Strict prefixes and lengths; known noisy patterns removed"],
        ["Network / domain / email / file hashes", "IPv4/IPv6/CIDR, URL/domain/onion, email, MD5/SHA", "Private/localhost IPs skipped; generic families run last"],
        ["Infrastructure / SaaS / shadow IT", "config/database files, public object storage, tunnels", "Context-bound expressions and minimum lengths"],
        ["Financial and identity", "card-shaped values, SSN, IBAN, SWIFT BIC", "Country-code-constrained case-sensitive BIC; no broad ACH matcher"],
        ["Actor / exfiltration / CVE / crypto", "ransomware/APT names, dump evidence, CVE IDs, addresses", "Evidence patterns; not customer attribution by themselves"],
    ], [2700, 3600, 3060])
    add_para(doc, "A match starts with the pattern's base confidence. If a supplied customer domain appears in the same line, confidence increases by 0.10 and is capped at 1.00.")

    doc.add_heading("PII normalization and graph entities", level=1)
    add_bullets(doc, [
        "Email: trim and lowercase; reject missing-@ and empty/nan values.",
        "Phone: keep digits, normalize common +91/0 forms in the graph path and common country/trunk forms in direct search.",
        "Name: collapse whitespace and normalize title case for graph nodes; lower/collapse for search input.",
        "Aadhaar-shaped input: detect 4-4-4 digits and remove spaces for lookup. Detection is formatting only, not checksum validation.",
        "Graph extraction produces EMAIL, PHONE, NAME, and ADDRESS entities and deduplicates them per record.",
    ])

    doc.add_heading("Customer attribution", level=1)
    add_bullets(doc, [
        "Exact IP and CIDR membership for owned network ranges.",
        "Exact-domain and subdomain-boundary matching; delimited path/text occurrences remain lower-confidence keyword evidence.",
        "Credential/email-domain matching for registered customer domains and executives.",
        "CVE-to-product matching through CveProductMap with normalized product names and conservative version-range evaluation.",
        "Brand/keyword matching in dark-web and ransomware evidence.",
        "Unattributable financial identifiers remain global threat signals; they are not falsely assigned to a customer.",
    ])

    doc.add_heading("STIX 2.1 mapping", level=1)
    add_table(doc, ["IOC", "STIX object/pattern"], [
        ["IPv4 / IPv6", "indicator with ipv4-addr / ipv6-addr value pattern"],
        ["Domain / URL / email", "indicator with domain-name / url / email-addr value pattern"],
        ["MD5 / SHA-1 / SHA-256 / SHA-512", "indicator with file:hashes pattern"],
        ["CVE", "vulnerability object with deterministic UUIDv5 identifier"],
        ["Non-standard Telegram/entity type", "indicator using open pattern_type 'argus', not invalid STIX syntax"],
    ], [2700, 6660])
    add_para(doc, "Telegram-derived identity, vulnerability, and indicator identifiers use UUIDv5 over a stable namespace/key; bundle IDs use UUIDv4. Values are escaped before insertion into STIX patterns, duplicate type/value pairs are removed, and confidence is clamped to 0–100.")

    add_pii_math(doc)

    add_comparison(doc, *PII_COMPARISON, title="PII and sensitive-data comparison")
    add_para(doc, "Selection context: Presidio is the strongest fit here for extensible self-hosted PII analysis and anonymization. Google Sensitive Data Protection provides managed inspection, discovery, custom detectors, likelihoods, and de-identification across cloud/hybrid workflows. Amazon Macie specializes in managed S3 sensitive-data discovery and findings. Argus is narrower as an enterprise DLP tool, but connects PII candidates to CTI indicators, customer assets, breach graphs, findings, and STIX.")
    add_comparison(doc, *IOC_STIX_COMPARISON, title="IOC extraction and CTI interoperability comparison")
    add_para(doc, "Selection context: iocextract is a focused library for extracting and refanging common or encoded IOCs. OpenCTI and MISP are mature CTI knowledge-sharing platforms with graph/event models and STIX-oriented interoperability. Argus combines broader credential/PII/IOC patterns with conservative customer matching, deterministic output identifiers, breach-graph linkage, and remediation workflow; it currently lacks iocextract's dedicated defang/refang coverage.")
    add_sources(doc, ["kleene", "stix", "rfc9562", "fips180", "presidio", "presidio_entities", "presidio_anonymizer", "google_dlp", "google_dlp_custom", "macie", "macie_custom", "iocextract", "opencti", "opencti_connectors", "misp"])
    add_verification(doc, ["Pattern ordering and false-positive regressions are tested, including rejection of an uppercase English word as a SWIFT BIC.", "STIX tests validate standard patterns, deterministic IDs, CVE vulnerability objects, escaping, confidence bounds, and non-standard pattern handling."])
    add_boundaries(doc, [
        "Regex matches are candidates, not proof of validity. Card numbers, national identifiers, and secrets require contextual or checksum/provider validation before action.",
        "This is CTI-focused PII/IOC correlation, not a general enterprise DLP discovery engine for every storage platform.",
        "Sensitive values may appear in evidence; access control, retention, encryption, and redaction policies remain operator responsibilities.",
    ])
    save_with_embedded_svg(doc, MODULES / "05_PII_IOC_and_STIX_Module.docx", "pii-ioc-stix-flow")


def walkthrough_document(output_path: Path | None = None):
    doc = new_document("Complete Walkthrough")
    masthead(doc, "Argus Unified — Complete System Walkthrough", "Architecture, module integration, deployment, operator flow, algorithms, evidence, and current limits")
    doc.add_heading("What was consolidated", level=1)
    add_para(doc, "Argus Unified is a new, clean project folder assembled around argus_main. No original source folder was modified. The standalone OSINT and breach codebase was merged once under backend/modules and exposed through the shared edge and UI. The unique Telegram monitor and PII-link behavior was reimplemented as backend/modules/telegram_intel and connected to the existing detection, correlation, finding, SLA, exposure, STIX, Redis/Celery, and customer-asset systems.")
    add_table(doc, ["Source", "Unified destination", "Duplication decision"], [
        ["argus_main", "Project shell, API, UI, infrastructure, scoring", "Canonical base retained"],
        ["osint and breach codebase", "backend/modules/osint and backend/modules/breach; modules-ui", "One mounted copy; duplicate product shell removed"],
        ["cyber_threat_monitor", "backend/modules/telegram_intel", "Rules/workflow retained; duplicate SQLite services removed"],
        ["pii_link", "Shared IOC/PII/STIX and Telegram correlation", "Concepts integrated; duplicate ES/Redis/Neo4j stack removed"],
    ], [2450, 3950, 2960])
    add_flow(doc, "unified-platform-flow", "Figure 1. Unified request, execution, storage, and analyst-output path.")

    doc.add_heading("System map", level=1)
    add_table(doc, ["Area", "Location", "Purpose"], [
        ["Core API", "backend/arguswatch", "FastAPI, models, collectors, engines, RBAC, metrics"],
        ["Integrated modules", "backend/modules", "OSINT, breach search/graph, Telegram intelligence"],
        ["Dashboard", "backend/arguswatch/static", "Canonical analyst UI including native Telegram page"],
        ["Modules UI", "modules-ui", "OSINT and breach workflows served below /modules"],
        ["Database bootstrap", "initdb", "PostgreSQL schema, migrations, Telegram tables"],
        ["Deployment", "docker/Dockerfile, docker/docker-compose.yml, setup/ launchers", "One Argus image reused across isolated runtime roles"],
        ["Tests", "backend/tests", "Unit, regression, API smoke, and integration specifications"],
        ["Documentation", "docs/modules and docs/flowcharts", "Verified per-module reports and monochrome flows"],
    ], [1700, 3250, 4410])

    doc.add_heading("Secure deployment walkthrough", level=1)
    add_numbered(doc, [
        "Run .\\setup\\setup.ps1 or setup\\setup.bat to create .env with independent high-entropy PostgreSQL, Redis, JWT, administrator, ClickHouse, and breach-hash secrets; then review optional provider settings.",
        "Keep the default 127.0.0.1 bind for local use. For remote production access, terminate TLS at an approved reverse proxy, load balancer, VPN, or cluster ingress and configure allowed origins and trusted proxy addresses.",
        "Start with .\\setup\\start.ps1 or setup\\start.bat. The Python startup module waits for PostgreSQL, runs the unified migration, and replaces itself with Uvicorn.",
        "Open http://127.0.0.1:7777. Create users/roles and onboard verified customer-owned assets before running matching or recon.",
        "If Telegram collection is required, set TELEGRAM_API_ID, TELEGRAM_API_HASH, TELEGRAM_PHONE and TELEGRAM_CHANNELS, then run the one-time authorization command documented in README.md.",
        "Run collectors or explicit imports within authorized scope. Confirm collector status, queue health, and datastore metrics before relying on coverage.",
        "Back up PostgreSQL, ClickHouse datasets, search indices, and the persistent Telegram session under the operator's secrets/retention policy.",
    ])
    add_code(doc, ".\\setup\\setup.ps1\n.\\setup\\start.ps1\n# Command Prompt alternatives: setup\\setup.bat, setup\\start.bat\n\n# Validate configuration without starting containers\ndocker compose --project-directory . -f docker/docker-compose.yml config")

    doc.add_heading("Analyst walkthrough", level=1)
    add_numbered(doc, [
        "Onboard a customer and its owned assets. The matching gate blocks a newly created customer until assets exist.",
        "Use OSINT for subject discovery. Treat confidence/evidence tiers as analyst aids and preserve source URLs and pivot lineage.",
        "Use breach search for direct identifiers and the bounded graph for relationship hypotheses. Review common-value warnings and provenance.",
        "Use Telegram Analyze for stateless triage, Import for authorized exports, or the collector for authorized configured channels. Review risk reasons and extracted IOCs.",
        "Run correlation. Direct evidence becomes customer findings; unmatched activity contributes to sector/global pressure rather than false attribution.",
        "Review severity, match proof, exposure dimensions, actor/campaign context, SLA, and remediation steps. Human approval remains required for consequential action.",
        "Export STIX/SIEM or reports after evidence review.",
    ])

    doc.add_heading("Implemented mathematical reference", level=1)
    add_table(doc, ["Module", "Method", "Implemented expression / rule"], [
        ["Core", "Exposure", "E=max(.50D1+.30D2+.20D3,.20D4); M=.75+.00125D4+.00125D5; Risk=min(100,E·M)"],
        ["Core", "Freshness", "decay=exp(−ln(2)·age/half_life); normalized=raw·feed_confidence·decay"],
        ["Core", "Pressure", "min(10,3.3log10(count)), recency boost, 7-day decay; sector=.60max+.40mean"],
        ["OSINT", "DOM similarity", "Σ min(tag counts) / Σ max(tag counts)"],
        ["OSINT", "Confidence tiers", "95 confirmed; 75 high; 50 medium; 30 ambiguous; 1 unverified; 0 not found"],
        ["Breach", "Graph", "Bounded BFS; >200 common-value short circuit; ≤500 records + ≤300 entities"],
        ["Telegram", "Category confidence", "c=1−exp(−raw_weight/4)"],
        ["Telegram", "Risk", "p=1−exp(−Σ impact·confidence), then complement boosts; flag at 55 or ransomware≥.65"],
        ["IOC/STIX", "Detection/export", "92 patterns / 17 ordered families; confidence cap; deterministic UUIDv5 STIX object IDs"],
    ], [1100, 1900, 6360], font_size=7.7)

    doc.add_heading("Complete module flowcharts and formal algorithms", level=1)
    add_para(doc, "This section contains every module flowchart used by the five authoritative module reports. Each diagram is immediately followed by the corresponding code-faithful mathematical specification, executable pseudocode, bounds, invariants, complexity statement, and research/standards basis. The opening unified flowchart is Figure 1; the five module flows below are Figures 2–6.")

    doc.add_heading("Core platform", level=2)
    add_flow(doc, "core-platform-flow", "Figure 2. Core ingestion, correlation, scoring, and response path.")
    add_core_math(doc, level=2)

    doc.add_heading("OSINT investigation", level=2)
    add_flow(doc, "osint-investigation-flow", "Figure 3. Profile-aware OSINT investigation and verification path.")
    add_osint_math(doc, level=2)

    doc.add_heading("Breach search and graph", level=2)
    add_flow(doc, "breach-correlation-flow", "Figure 4. Direct breach lookup, correlation, caching, and graph expansion.")
    add_breach_math(doc, level=2)

    doc.add_heading("Telegram threat intelligence", level=2)
    add_flow(doc, "telegram-intelligence-flow", "Figure 5. Telegram ingestion, classification, correlation, and response path.")
    add_telegram_math(doc, level=2)

    doc.add_heading("PII, IOC correlation and STIX", level=2)
    add_flow(doc, "pii-ioc-stix-flow", "Figure 6. Shared PII/IOC extraction, customer correlation, and STIX mapping.")
    add_pii_math(doc, level=2)

    add_unified_math(doc)

    doc.add_heading("Consolidated market alternatives", level=1)
    add_para(doc, "The following matrices consolidate the module-level research into the system walkthrough. They compare native, evidenced capability rather than marketing similarity. Argus checks reflect the present repository only; external checks reflect the cited official product documentation reviewed on 10 August 2026. A product can often gain crossed capabilities through custom development, paid add-ons, connectors, or partner data.")
    add_table(doc, ["Argus area", "Representative alternatives", "Where the alternative is strongest", "Argus-specific integration"], [
        ["Core CTI platform", "OpenCTI, MISP, ThreatConnect, Recorded Future", "Mature CTI sharing/graphs, commercial case tooling, or managed intelligence corpus", "Owned-asset correlation plus recon, breach, Telegram, exposure, SLA, and remediation"],
        ["OSINT investigation", "SpiderFoot, Maltego, Sherlock, Maigret", "Broad modules, graph analysis/collaboration, or focused username enumeration", "Subject disambiguation, evidence tiers, calibration, resumable jobs, downstream pivots"],
        ["Breach search", "HIBP, Intelligence X, LeakCheck, DeHashed", "Hosted corpora, notifications, and externally maintained search coverage", "Operator-controlled stores, bounded PII graph, stable IDs, customer finding handoff"],
        ["Telegram intelligence", "Telethon, Maltego Monitor, Social Links, OpenCTI", "Raw MTProto development, commercial monitoring/transforms, or general CTI graphing", "Deterministic cybercrime scoring, bounded import, IOC/customer matching, STIX"],
        ["PII / IOC / STIX", "Presidio, Google SDP, Macie, iocextract, OpenCTI, MISP", "PII anonymization, managed data discovery, defanged IOC extraction, CTI sharing", "One ordered detector feeding attribution, breach graph, findings, and deterministic STIX"],
    ], [1500, 2400, 2780, 2680], font_size=7.2)
    add_comparison(doc, *CORE_COMPARISON, title="Core platform alternatives")
    add_comparison(doc, *OSINT_COMPARISON, title="OSINT alternatives")
    add_comparison(doc, *BREACH_COMPARISON, title="Breach-search alternatives")
    add_comparison(doc, *TELEGRAM_COMPARISON, title="Telegram-intelligence alternatives")
    add_comparison(doc, *PII_COMPARISON, title="PII and sensitive-data alternatives")
    add_comparison(doc, *IOC_STIX_COMPARISON, title="IOC extraction and STIX alternatives")

    doc.add_heading("Configuration and migration", level=1)
    add_bullets(doc, [
        "The unified migration is backend/arguswatch/scripts/migrate_unified.py; database bootstrap includes initdb/15_telegram_intel.sql.",
        "Auth and AI defaults are secure: AUTH_DISABLED=false and AI_AUTONOMY_ENABLED=false unless deliberately changed.",
        "Only the Argus API port is published and it binds to 127.0.0.1 by default; infrastructure and internal application roles publish no host ports.",
        "The canonical docker/Dockerfile builds the dashboard, modules SPA, API, workers, breach roles, intelligence proxy, and recon tools into the same non-root Argus image.",
        "Python and JavaScript dependency pins were updated and audited; deprecated Starlette WSGI middleware was replaced with a2wsgi.",
        "The multi-stage image builds Python wheels and the modules SPA once, pins recon-tool releases, runs as a non-root user, and keeps test dependencies outside the production image.",
    ])

    doc.add_heading("Verification ledger", level=1)
    add_table(doc, ["Check", "Result", "Boundary"], [
        ["Python parse", "193 files passed", "Syntax/AST only"],
        ["Backend unit/stateless suite", "242 passed", "DB-heavy suites excluded from this pass"],
        ["Application smoke", "9 passed", "Boot/CORS/validation/auth without live data stores"],
        ["Frontend lint/build", "Passed", "Optimized Vite output generated"],
        ["Browser input", "Passed", "Telegram analysis verified at 1600×1000"],
        ["Python dependency audit", "0 known vulnerabilities", "Based on current advisory database"],
        ["npm audit", "0 known vulnerabilities", "Current lockfile"],
        ["Compose resolution", "Passed", "Configuration only"],
        ["Live Docker stack", "Not run", "Local Docker Linux engine unavailable"],
        ["Live Telegram collection", "Not run", "No authorized credentials/session supplied"],
    ], [3100, 2350, 3910])

    doc.add_heading("Documentation set", level=1)
    add_bullets(doc, [
        "01_Argus_Core_Platform.docx",
        "02_OSINT_Investigation_Module.docx",
        "03_Breach_Search_and_Graph_Module.docx",
        "04_Telegram_Threat_Intelligence_Module.docx",
        "05_PII_IOC_and_STIX_Module.docx",
        "Six matching black-and-white SVG flowcharts in docs/flowcharts.",
    ])
    add_para(doc, "These reports are the authoritative implemented-state documentation for this consolidation. The source DOCX files remain historical input material.")
    add_boundaries(doc, [
        "Production readiness is conditional on environment-specific deployment validation, approved secrets/TLS, backup/restore testing, capacity testing, and legal authorization for every data source.",
        "No document asserts that Argus universally outperforms commercial platforms. The comparisons expose native feature differences and the unified workflow's specific strengths and gaps.",
        "External services can change without code changes; revalidate provider behavior, quotas, and contracts before each release.",
    ])
    doc.add_heading("Source material and market references", level=1)
    add_para(doc, "Existing project documents reused as implementation lineage:")
    add_bullets(doc, SOURCE_DOCS)
    add_para(doc, "Official sources used for comparison matrices:")
    for key in SOURCES:
        label, url = SOURCES[key]
        p = doc.add_paragraph(style="List Bullet")
        _hyperlink(p, label, url)
    save_with_embedded_svgs(doc, output_path or (DOCS / "Argus_Unified_Complete_Walkthrough.docx"), [
        "unified-platform-flow",
        "core-platform-flow",
        "osint-investigation-flow",
        "breach-correlation-flow",
        "telegram-intelligence-flow",
        "pii-ioc-stix-flow",
    ])


def build_index():
    index = """# Argus Unified documentation

These reports describe the consolidated implementation as verified on 10 August 2026. They intentionally exclude roadmap items and unimplemented claims from the source-project documents.

## Module reports

- [Core Platform](modules/01_Argus_Core_Platform.docx)
- [OSINT Investigation](modules/02_OSINT_Investigation_Module.docx)
- [Breach Search and Graph](modules/03_Breach_Search_and_Graph_Module.docx)
- [Telegram Threat Intelligence](modules/04_Telegram_Threat_Intelligence_Module.docx)
- [PII, IOC and STIX](modules/05_PII_IOC_and_STIX_Module.docx)
- [Complete System Walkthrough](Argus_Unified_Complete_Walkthrough.docx)

## Black-and-white SVG flowcharts

- [Core platform flow](flowcharts/core-platform-flow.svg)
- [OSINT investigation flow](flowcharts/osint-investigation-flow.svg)
- [Breach correlation flow](flowcharts/breach-correlation-flow.svg)
- [Telegram intelligence flow](flowcharts/telegram-intelligence-flow.svg)
- [PII / IOC / STIX flow](flowcharts/pii-ioc-stix-flow.svg)
- [Unified platform flow](flowcharts/unified-platform-flow.svg)

## Verification boundary

The isolated Python suite, application smoke tests, frontend lint/build, dependency audits, Compose resolution, and browser input flow passed. A full live multi-database Docker pass still requires a host with a running Docker Linux engine, and live Telegram collection requires an authorized user session.
"""
    (DOCS / "README.md").write_text(index, encoding="utf-8")


def main():
    build_flows()
    core_document()
    osint_document()
    breach_document()
    telegram_document()
    pii_document()
    walkthrough_document()
    build_index()
    print("Generated 6 DOCX reports, 6 SVG flowcharts, 6 PNG previews, and docs/README.md")


if __name__ == "__main__":
    main()
